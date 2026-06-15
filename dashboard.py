# ==========================================================
# DASHBOARD.PY
# DASHBOARD / HISTORIAL DE EVENTOS - VOLCAN APP
# MEJORA: PDF/WORD VERTICAL + GRAFICAS POR GRUPO OPERATIVO + TABLAS AMPLIADAS + FALLA/CAUSA
# ==========================================================

import streamlit as st
import pandas as pd
import altair as alt
import base64

from io import BytesIO
from datetime import datetime

import matplotlib.pyplot as plt
from PIL import Image as PILImage

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    KeepTogether,
    Image as RLImage
)

from database import cargar_bitacora, cargar_equipos


# ==========================================================
# UTILIDADES GENERALES
# ==========================================================

def obtener_bytes_imagen(foto_base64):

    if not isinstance(foto_base64, str):
        return None

    if foto_base64.strip() in ["", "SIN FOTO"]:
        return None

    if "base64," in foto_base64:
        foto_base64 = foto_base64.split("base64,")[1]

    try:
        return base64.b64decode(foto_base64)
    except Exception:
        return None


def normalizar_texto(valor):
    return str(valor).strip().upper()


def texto_estado_operativo_resumen(valor):
    texto = str(valor).strip()
    if not texto:
        return ""
    texto = texto.replace("OPERATIVAS:", "OP:")
    texto = texto.replace("OPERATIVA:", "OP:")
    texto = texto.replace("STAND BY:", "SB:")
    return texto


def texto_condicion_resumen(valor):
    texto = str(valor).strip()
    reemplazos = {
        "GLOBAL GRUPO OPERATIVO": "GLOBAL GRUPO",
        "GLOBAL NIVEL OPERATIVO": "GLOBAL NIVEL",
        "GLOBAL SISTEMA OPERATIVO": "GLOBAL SISTEMA",
    }
    return reemplazos.get(texto.upper(), texto)


def obtener_grupo_operativo(nivel, ubicacion, codigo):
    """
    Regla operacional:
    - En niveles con trenes, si se interviene una bomba del tren, se considera afectado todo el tren.
    - En niveles sin tren, se considera afectada solo la bomba intervenida.
    """

    nivel_txt = normalizar_texto(nivel)
    ubicacion_txt = normalizar_texto(ubicacion)
    codigo_txt = normalizar_texto(codigo)

    if "TREN" in ubicacion_txt:
        if "PRIMER" in ubicacion_txt:
            return f"{nivel_txt} | PRIMER TREN"
        if "SEGUNDO" in ubicacion_txt:
            return f"{nivel_txt} | SEGUNDO TREN"
        return f"{nivel_txt} | TREN"

    return f"{nivel_txt} | {codigo_txt}"


def obtener_estado_operativo_bomba(nivel, ubicacion, codigo=""):
    """
    Clasificación operacional definida para el sistema de bombeo.

    Reglas actuales:
    - NIVEL 10: no tiene bombas en stand by.
    - NIVEL 5 / POZO 2: POSICION 01 y POSICION 02 están en STAND BY.
    - NIVEL 11: POSICION 02 está en STAND BY.

    Nota técnica:
    Las bombas en STAND BY se muestran en tablas y gráficas, pero no deben castigar
    la disponibilidad mecánica operativa del grupo, nivel o sistema.
    """

    nivel_txt = normalizar_texto(nivel)
    ubicacion_txt = normalizar_texto(ubicacion)

    if nivel_txt == "NIVEL 5" and ubicacion_txt in ["POSICION 01", "POSICION 02"]:
        return "STAND BY"

    if nivel_txt == "NIVEL 11" and ubicacion_txt == "POSICION 02":
        return "STAND BY"

    return "OPERATIVA"


def es_bomba_operativa(valor_estado):
    return normalizar_texto(valor_estado) == "OPERATIVA"


def preparar_equipos_detalle():

    df_equipos = cargar_equipos()

    if df_equipos.empty:
        return pd.DataFrame(columns=["sistema", "nivel", "ubicacion", "codigo", "grupo_operativo", "estado_operativo"])

    df_equipos.columns = df_equipos.columns.str.strip().str.lower()

    columnas = ["sistema", "nivel", "ubicacion", "codigo"]

    for col in columnas:
        if col not in df_equipos.columns:
            df_equipos[col] = ""

    for col in columnas:
        df_equipos[col] = df_equipos[col].astype(str).str.strip()

    df_equipos = df_equipos.drop_duplicates(subset=["nivel", "ubicacion", "codigo"])

    df_equipos["grupo_operativo"] = df_equipos.apply(
        lambda row: obtener_grupo_operativo(row["nivel"], row["ubicacion"], row["codigo"]),
        axis=1
    )

    df_equipos["estado_operativo"] = df_equipos.apply(
        lambda row: obtener_estado_operativo_bomba(row["nivel"], row["ubicacion"], row["codigo"]),
        axis=1
    )

    return df_equipos


def preparar_equipos_por_nivel():

    df_equipos = preparar_equipos_detalle()

    if df_equipos.empty:
        return {}

    return (
        df_equipos
        .drop_duplicates(subset=["nivel", "codigo"])
        .groupby("nivel")["codigo"]
        .nunique()
        .to_dict()
    )


def crear_resumen_diario_nivel(df_eventos, df_equipos, nivel, fecha_inicio, fecha_fin):
    """
    Resumen diario del nivel completo.
    - Fallas = eventos reales registrados en el nivel, no se multiplican por bomba.
    - Horas reales = duración real de los eventos.
    - Horas-equipo operativas = horas reales x cantidad de bombas OPERATIVAS afectadas.
    - Las bombas STAND BY se muestran, pero no castigan la disponibilidad mecánica operativa.
    - Disponibilidad del nivel = (horas calendario operativo - horas-equipo operativas) / horas calendario operativo.
    """

    df_eq_nv = df_equipos[
        df_equipos["nivel"].astype(str).str.strip() == str(nivel).strip()
    ].copy()

    fechas = pd.date_range(fecha_inicio, fecha_fin, freq="D").date
    registros = []

    if df_eq_nv.empty:
        return pd.DataFrame(registros)

    if "estado_operativo" not in df_eq_nv.columns:
        df_eq_nv["estado_operativo"] = df_eq_nv.apply(
            lambda row: obtener_estado_operativo_bomba(row.get("nivel", ""), row.get("ubicacion", ""), row.get("codigo", "")),
            axis=1
        )

    df_eq_nv_operativas = df_eq_nv[df_eq_nv["estado_operativo"].astype(str).str.upper() == "OPERATIVA"].copy()

    cantidad_bombas_nivel = df_eq_nv["codigo"].nunique()
    cantidad_bombas_operativas = df_eq_nv_operativas["codigo"].nunique()
    cantidad_bombas_standby = max(cantidad_bombas_nivel - cantidad_bombas_operativas, 0)
    horas_calendario_dia = cantidad_bombas_operativas * 24

    df_ev = df_eventos.copy()
    if not df_ev.empty:
        df_ev["fecha_dia"] = df_ev["fecha"].dt.date
        df_ev["grupo_operativo"] = df_ev.apply(
            lambda row: obtener_grupo_operativo(row.get("nivel", ""), row.get("ubicacion", ""), row.get("codigo", "")),
            axis=1
        )

    for fecha in fechas:
        df_dia = df_ev[
            (df_ev["nivel"].astype(str).str.strip() == str(nivel).strip()) &
            (df_ev["fecha_dia"] == fecha)
        ].copy() if not df_ev.empty else pd.DataFrame()

        horas_reales = 0.0
        horas_equipo = 0.0
        cantidad_fallas = 0
        bombas_afectadas = []
        bombas_intervenidas = []
        bombas_standby = df_eq_nv[df_eq_nv["estado_operativo"].astype(str).str.upper() == "STAND BY"]["codigo"].astype(str).tolist()

        if not df_dia.empty:
            cantidad_fallas = int(df_dia.shape[0])
            horas_reales = float(df_dia["tiempo_parada"].sum())

            for _, ev in df_dia.iterrows():
                grupo = str(ev.get("grupo_operativo", ""))
                horas_evento = float(ev.get("tiempo_parada", 0) or 0)

                df_afectadas = df_eq_nv[df_eq_nv["grupo_operativo"] == grupo]
                df_afectadas_operativas = df_afectadas[
                    df_afectadas["estado_operativo"].astype(str).str.upper() == "OPERATIVA"
                ]
                cantidad_afectadas_operativas = df_afectadas_operativas["codigo"].nunique()

                horas_equipo += horas_evento * cantidad_afectadas_operativas
                bombas_afectadas += df_afectadas_operativas["codigo"].astype(str).tolist()
                bombas_intervenidas.append(str(ev.get("codigo", "")))

        disponibilidad = (
            ((horas_calendario_dia - horas_equipo) / horas_calendario_dia) * 100
            if horas_calendario_dia > 0
            else 100
        )
        disponibilidad = max(0, min(disponibilidad, 100))

        registros.append({
            "fecha": fecha,
            "nivel": nivel,
            "bombas_nivel": cantidad_bombas_nivel,
            "bombas_operativas": cantidad_bombas_operativas,
            "bombas_standby": cantidad_bombas_standby,
            "bombas_intervenidas": ", ".join(sorted(set([b for b in bombas_intervenidas if b.strip() != ""]))),
            "bombas_afectadas": ", ".join(sorted(set([b for b in bombas_afectadas if b.strip() != ""]))),
            "bombas_standby_codigos": ", ".join(sorted(set([b for b in bombas_standby if b.strip() != ""]))),
            "horas_reales": round(horas_reales, 2),
            "horas_equipo": round(horas_equipo, 2),
            "cantidad_fallas": cantidad_fallas,
            "disponibilidad": round(disponibilidad, 2)
        })

    return pd.DataFrame(registros)

def crear_matriz_disponibilidad_nivel(df_eventos, df_equipos, nivel, fecha_inicio, fecha_fin):
    """
    Matriz diaria por bomba.
    - OPERATIVA: entra al cálculo de disponibilidad.
    - STAND BY: se muestra en tabla/gráfica, pero no castiga disponibilidad operativa.
    - INTERVENIDO: bomba donde se registró directamente el evento.
    - AFECTADO POR TREN: bomba operativa que no fue intervenida, pero pertenece al mismo tren/grupo operativo detenido.
    - N/I: bomba operativa sin intervención ni afectación.
    """

    df_eq_nv = df_equipos[
        df_equipos["nivel"].astype(str).str.strip() == str(nivel).strip()
    ].copy()

    fechas = pd.date_range(fecha_inicio, fecha_fin, freq="D").date
    registros = []

    if df_eq_nv.empty:
        return pd.DataFrame(registros)

    if "estado_operativo" not in df_eq_nv.columns:
        df_eq_nv["estado_operativo"] = df_eq_nv.apply(
            lambda row: obtener_estado_operativo_bomba(row.get("nivel", ""), row.get("ubicacion", ""), row.get("codigo", "")),
            axis=1
        )

    df_ev = df_eventos.copy()
    if not df_ev.empty:
        df_ev["fecha_dia"] = df_ev["fecha"].dt.date
        df_ev["grupo_operativo"] = df_ev.apply(
            lambda row: obtener_grupo_operativo(row.get("nivel", ""), row.get("ubicacion", ""), row.get("codigo", "")),
            axis=1
        )

    for fecha in fechas:
        df_dia = df_ev[
            (df_ev["nivel"].astype(str).str.strip() == str(nivel).strip()) &
            (df_ev["fecha_dia"] == fecha)
        ].copy() if not df_ev.empty else pd.DataFrame()

        for _, eq in df_eq_nv.iterrows():
            codigo = str(eq["codigo"])
            ubicacion = str(eq["ubicacion"])
            grupo = str(eq["grupo_operativo"])
            estado_operativo = str(eq.get("estado_operativo", obtener_estado_operativo_bomba(nivel, ubicacion, codigo)))
            es_standby = normalizar_texto(estado_operativo) == "STAND BY"

            eventos_directos = pd.DataFrame()
            eventos_afectan = pd.DataFrame()

            if not df_dia.empty:
                eventos_directos = df_dia[df_dia["codigo"].astype(str).str.strip() == codigo.strip()]
                eventos_afectan = df_dia[df_dia["grupo_operativo"] == grupo]

            horas_reales_intervencion = float(eventos_directos["tiempo_parada"].sum()) if not eventos_directos.empty else 0.0
            horas_afectadas = float(eventos_afectan["tiempo_parada"].sum()) if not eventos_afectan.empty else 0.0
            fallas_directas = int(eventos_directos.shape[0]) if not eventos_directos.empty else 0
            fallas_afectan = int(eventos_afectan.shape[0]) if not eventos_afectan.empty else 0

            if not eventos_directos.empty:
                tipos_falla = sorted(set([
                    str(x).strip()
                    for x in eventos_directos.get("tipo_falla", pd.Series(dtype=str)).tolist()
                    if str(x).strip() not in ["", "nan", "None"]
                ]))
                causas_preliminares = sorted(set([
                    str(x).strip()
                    for x in eventos_directos.get("causa_preliminar", pd.Series(dtype=str)).tolist()
                    if str(x).strip() not in ["", "nan", "None"]
                ]))
                tipo_falla_resumen = " | ".join(tipos_falla) if tipos_falla else "N/D"
                causa_preliminar_resumen = " | ".join(causas_preliminares) if causas_preliminares else "N/D"
            else:
                tipo_falla_resumen = "N/I"
                causa_preliminar_resumen = "N/I"

            if es_standby:
                # La bomba en reserva se informa, pero no reduce disponibilidad operativa.
                disponibilidad = 100.0
                if fallas_directas > 0 or horas_reales_intervencion > 0:
                    condicion = "INTERVENIDO - STAND BY"
                else:
                    condicion = "STAND BY"
            else:
                disponibilidad = ((24 - horas_afectadas) / 24) * 100
                disponibilidad = max(0, min(disponibilidad, 100))

                if fallas_directas > 0 or horas_reales_intervencion > 0:
                    condicion = "INTERVENIDO"
                elif fallas_afectan > 0 or horas_afectadas > 0:
                    condicion = "AFECTADO POR TREN"
                else:
                    condicion = "N/I"

            registros.append({
                "fecha": fecha,
                "nivel": nivel,
                "ubicacion": ubicacion,
                "codigo": codigo,
                "grupo_operativo": grupo,
                "estado_operativo": estado_operativo,
                "horas_reales_intervencion": round(horas_reales_intervencion, 2),
                "horas_afectadas": round(horas_afectadas, 2),
                "cantidad_fallas_directas": fallas_directas,
                "cantidad_fallas_afectan": fallas_afectan,
                "tipo_falla": tipo_falla_resumen,
                "causa_preliminar": causa_preliminar_resumen,
                "disponibilidad": round(disponibilidad, 2),
                "intervencion": condicion
            })

    return pd.DataFrame(registros)

def crear_resumen_nivel(df_matriz):

    if df_matriz.empty:
        return pd.DataFrame()

    df_resumen = (
        df_matriz
        .groupby(["nivel", "ubicacion", "codigo", "grupo_operativo", "estado_operativo"], as_index=False)
        .agg(
            horas_reales_intervencion=("horas_reales_intervencion", "sum"),
            horas_afectadas=("horas_afectadas", "sum"),
            cantidad_fallas_directas=("cantidad_fallas_directas", "sum"),
            cantidad_fallas_afectan=("cantidad_fallas_afectan", "sum"),
            disponibilidad_promedio=("disponibilidad", "mean")
        )
    )

    df_resumen["disponibilidad_promedio"] = df_resumen["disponibilidad_promedio"].round(2)
    df_resumen["horas_reales_intervencion"] = df_resumen["horas_reales_intervencion"].round(2)
    df_resumen["horas_afectadas"] = df_resumen["horas_afectadas"].round(2)

    def condicion_periodo(row):
        if normalizar_texto(row.get("estado_operativo", "")) == "STAND BY":
            if row["cantidad_fallas_directas"] > 0 or row["horas_reales_intervencion"] > 0:
                return "INTERVENIDO - STAND BY"
            return "STAND BY"
        if row["cantidad_fallas_directas"] > 0 or row["horas_reales_intervencion"] > 0:
            return "INTERVENIDO"
        if row["cantidad_fallas_afectan"] > 0 or row["horas_afectadas"] > 0:
            return "AFECTADO POR TREN"
        return "N/I"

    df_resumen["intervencion"] = df_resumen.apply(condicion_periodo, axis=1)

    return df_resumen


def crear_detalle_diario_bombas(df_matriz):

    if df_matriz.empty:
        return pd.DataFrame()

    df_detalle = df_matriz.copy()
    df_detalle = df_detalle.sort_values(["fecha", "ubicacion", "codigo"])
    return df_detalle


def crear_resumen_general_disponibilidad(df_eventos, df_equipos, niveles, fecha_inicio, fecha_fin):
    """
    Tabla general consolidada de disponibilidad.
    Estructura:
    - Detalle por bomba.
    - Global por grupo operativo / tren / pozo.
    - Global por nivel.
    - Global sistema.

    Regla STAND BY:
    - Las bombas STAND BY se reportan, pero no castigan la disponibilidad operativa.
    - Los globales se calculan con bombas OPERATIVAS.
    """

    registros = []
    matrices_globales = []

    for nv in niveles:

        df_eq_nv = df_equipos[
            df_equipos["nivel"].astype(str).str.strip() == str(nv).strip()
        ].copy()

        if df_eq_nv.empty:
            continue

        df_matriz_nv = crear_matriz_disponibilidad_nivel(
            df_eventos,
            df_equipos,
            nv,
            fecha_inicio,
            fecha_fin
        )

        if df_matriz_nv.empty:
            continue

        matrices_globales.append(df_matriz_nv.copy())

        subgrupos = obtener_subgrupos_nivel(df_eq_nv, nv)

        for subgrupo in subgrupos:

            df_matriz_sub = filtrar_matriz_subgrupo(
                df_matriz_nv,
                subgrupo["codigos"]
            )

            if df_matriz_sub.empty:
                continue

            # -------------------------------
            # Detalle global por bomba
            # -------------------------------
            df_bombas = (
                df_matriz_sub
                .groupby(["nivel", "ubicacion", "codigo", "estado_operativo"], as_index=False)
                .agg(
                    horas_reales=("horas_reales_intervencion", "sum"),
                    horas_afectadas=("horas_afectadas", "sum"),
                    fallas_directas=("cantidad_fallas_directas", "sum"),
                    fallas_afectan=("cantidad_fallas_afectan", "sum"),
                    disponibilidad=("disponibilidad", "mean")
                )
                .sort_values(["ubicacion", "codigo"])
            )

            for _, row in df_bombas.iterrows():

                if normalizar_texto(row.get("estado_operativo", "")) == "STAND BY":
                    if row["fallas_directas"] > 0 or row["horas_reales"] > 0:
                        condicion = "INTERVENIDO - STAND BY"
                    else:
                        condicion = "STAND BY"
                elif row["fallas_directas"] > 0 or row["horas_reales"] > 0:
                    condicion = "INTERVENIDO"
                elif row["fallas_afectan"] > 0 or row["horas_afectadas"] > 0:
                    condicion = "AFECTADO"
                else:
                    condicion = "N/I"

                registros.append({
                    "tipo": "BOMBA",
                    "nivel": str(row["nivel"]),
                    "grupo": subgrupo["nombre"],
                    "ubicacion": str(row["ubicacion"]),
                    "codigo": str(row["codigo"]),
                    "estado_operativo": str(row["estado_operativo"]),
                    "horas_reales": round(float(row["horas_reales"]), 2),
                    "horas_afectadas": round(float(row["horas_afectadas"]), 2),
                    "fallas": int(row["fallas_directas"]),
                    "disponibilidad": round(float(row["disponibilidad"]), 2),
                    "condicion": condicion
                })

            # -------------------------------
            # Global del subgrupo, solo bombas operativas
            # -------------------------------
            df_matriz_sub_operativas = df_matriz_sub[
                df_matriz_sub["estado_operativo"].astype(str).str.upper() == "OPERATIVA"
            ].copy()
            df_matriz_sub_standby = df_matriz_sub[
                df_matriz_sub["estado_operativo"].astype(str).str.upper() == "STAND BY"
            ].copy()

            registros.append({
                "tipo": "GLOBAL GRUPO",
                "nivel": str(nv),
                "grupo": subgrupo["nombre"],
                "ubicacion": "GLOBAL " + subgrupo["nombre"],
                "codigo": "-",
                "estado_operativo": f"OPERATIVAS: {df_matriz_sub_operativas['codigo'].nunique()} | STAND BY: {df_matriz_sub_standby['codigo'].nunique()}",
                "horas_reales": round(float(df_matriz_sub["horas_reales_intervencion"].sum()), 2),
                "horas_afectadas": round(float(df_matriz_sub_operativas["horas_afectadas"].sum()), 2),
                "fallas": int(df_matriz_sub["cantidad_fallas_directas"].sum()),
                "disponibilidad": round(float(df_matriz_sub_operativas["disponibilidad"].mean()), 2) if not df_matriz_sub_operativas.empty else 100.0,
                "condicion": "GLOBAL GRUPO OPERATIVO"
            })

        # -------------------------------
        # Global del nivel, solo bombas operativas
        # -------------------------------
        df_matriz_nv_operativas = df_matriz_nv[
            df_matriz_nv["estado_operativo"].astype(str).str.upper() == "OPERATIVA"
        ].copy()
        df_matriz_nv_standby = df_matriz_nv[
            df_matriz_nv["estado_operativo"].astype(str).str.upper() == "STAND BY"
        ].copy()

        registros.append({
            "tipo": "GLOBAL NIVEL",
            "nivel": str(nv),
            "grupo": "GLOBAL NIVEL",
            "ubicacion": "GLOBAL NIVEL",
            "codigo": "-",
            "estado_operativo": f"OPERATIVAS: {df_matriz_nv_operativas['codigo'].nunique()} | STAND BY: {df_matriz_nv_standby['codigo'].nunique()}",
            "horas_reales": round(float(df_matriz_nv["horas_reales_intervencion"].sum()), 2),
            "horas_afectadas": round(float(df_matriz_nv_operativas["horas_afectadas"].sum()), 2),
            "fallas": int(df_matriz_nv["cantidad_fallas_directas"].sum()),
            "disponibilidad": round(float(df_matriz_nv_operativas["disponibilidad"].mean()), 2) if not df_matriz_nv_operativas.empty else 100.0,
            "condicion": "GLOBAL NIVEL OPERATIVO"
        })

    # -------------------------------
    # Global del sistema completo, solo bombas operativas
    # -------------------------------
    if matrices_globales:
        df_global = pd.concat(matrices_globales, ignore_index=True)
        df_global_operativas = df_global[
            df_global["estado_operativo"].astype(str).str.upper() == "OPERATIVA"
        ].copy()
        df_global_standby = df_global[
            df_global["estado_operativo"].astype(str).str.upper() == "STAND BY"
        ].copy()

        registros.append({
            "tipo": "GLOBAL SISTEMA",
            "nivel": "TODOS",
            "grupo": "GLOBAL SISTEMA",
            "ubicacion": "GLOBAL SISTEMA",
            "codigo": "-",
            "estado_operativo": f"OPERATIVAS: {df_global_operativas['codigo'].nunique()} | STAND BY: {df_global_standby['codigo'].nunique()}",
            "horas_reales": round(float(df_global["horas_reales_intervencion"].sum()), 2),
            "horas_afectadas": round(float(df_global_operativas["horas_afectadas"].sum()), 2),
            "fallas": int(df_global["cantidad_fallas_directas"].sum()),
            "disponibilidad": round(float(df_global_operativas["disponibilidad"].mean()), 2) if not df_global_operativas.empty else 100.0,
            "condicion": "GLOBAL SISTEMA OPERATIVO"
        })

    return pd.DataFrame(registros)


# ==========================================================
# PDF - TABLAS Y GRAFICAS
# ==========================================================

def crear_tabla_pdf(data, headers, font_size=7, col_widths=None):

    tabla_data = [headers] + data

    tabla = Table(
        tabla_data,
        repeatRows=1,
        colWidths=col_widths
    )

    estilo = [
        # Encabezado principal
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#B71C1C")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),

        # Cuerpo de tabla
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), font_size),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [
            colors.white,
            colors.HexColor("#F5F5F5")
        ]),
    ]

    # ======================================================
    # RESALTADO GERENCIAL DE FILAS GLOBALES
    # ======================================================
    # Aplica a tablas donde la primera columna es "Tipo".
    # GLOBAL GRUPO   -> celeste suave
    # GLOBAL NIVEL   -> azul suave
    # GLOBAL SISTEMA -> verde suave
    for i, fila in enumerate(data, start=1):

        if not fila:
            continue

        texto_tipo = str(fila[0]).strip().upper()

        if texto_tipo == "GLOBAL GRUPO":
            estilo.extend([
                ("BACKGROUND", (0, i), (-1, i), colors.HexColor("#D9EDF7")),
                ("TEXTCOLOR", (0, i), (-1, i), colors.HexColor("#0C5460")),
                ("FONTNAME", (0, i), (-1, i), "Helvetica-Bold"),
            ])

        elif texto_tipo == "GLOBAL NIVEL":
            estilo.extend([
                ("BACKGROUND", (0, i), (-1, i), colors.HexColor("#B8DAFF")),
                ("TEXTCOLOR", (0, i), (-1, i), colors.HexColor("#003366")),
                ("FONTNAME", (0, i), (-1, i), "Helvetica-Bold"),
            ])

        elif texto_tipo == "GLOBAL SISTEMA":
            estilo.extend([
                ("BACKGROUND", (0, i), (-1, i), colors.HexColor("#D4EDDA")),
                ("TEXTCOLOR", (0, i), (-1, i), colors.HexColor("#155724")),
                ("FONTNAME", (0, i), (-1, i), "Helvetica-Bold"),
            ])

    tabla.setStyle(TableStyle(estilo))

    return tabla


def crear_imagen_pdf_desde_bytes(imagen_bytes, max_width=11.5 * cm, max_height=6.8 * cm):

    try:
        img_buffer = BytesIO(imagen_bytes)

        with PILImage.open(BytesIO(imagen_bytes)) as img:
            ancho_px, alto_px = img.size

        if ancho_px <= 0 or alto_px <= 0:
            return None

        escala = min(max_width / ancho_px, max_height / alto_px)

        imagen_pdf = RLImage(
            img_buffer,
            width=ancho_px * escala,
            height=alto_px * escala
        )
        imagen_pdf.hAlign = "CENTER"

        return imagen_pdf

    except Exception:
        return None


def obtener_subgrupos_nivel(df_eq_nv, nivel):
    """
    Divide cada nivel en grupos operativos para graficar.

    Reglas solicitadas:
    - NIVEL 10: dos gráficas, PRIMER TREN y SEGUNDO TREN.
    - NIVEL 11: una gráfica con POSICION 01, 02 y 03.
    - NIVEL 5: dos gráficas, POZO 1 = POSICION 04 y 05; POZO 2 = POSICION 01, 02 y 03.
    - Otros niveles: una gráfica con todas las bombas del nivel.
    """

    if df_eq_nv.empty:
        return []

    df_eq = df_eq_nv.copy()
    df_eq["nivel_norm"] = df_eq["nivel"].astype(str).str.upper().str.strip()
    df_eq["ubicacion_norm"] = df_eq["ubicacion"].astype(str).str.upper().str.strip()
    nivel_norm = str(nivel).upper().strip()

    subgrupos = []

    def agregar_subgrupo(nombre, filtro):
        df_sub = df_eq[filtro].copy()
        if df_sub.empty:
            return
        df_sub = df_sub.sort_values(["ubicacion", "codigo"])
        subgrupos.append({
            "nombre": nombre,
            "codigos": df_sub["codigo"].astype(str).tolist(),
            "df_equipos": df_sub
        })

    if "NIVEL 10" in nivel_norm:
        agregar_subgrupo(
            "PRIMER TREN",
            df_eq["ubicacion_norm"].str.contains("PRIMER TREN", na=False)
        )
        agregar_subgrupo(
            "SEGUNDO TREN",
            df_eq["ubicacion_norm"].str.contains("SEGUNDO TREN", na=False)
        )

    elif "NIVEL 11" in nivel_norm:
        agregar_subgrupo(
            "POSICIONES",
            df_eq["ubicacion_norm"].str.contains("POSICION", na=False)
        )

    elif "NIVEL 5" in nivel_norm:
        agregar_subgrupo(
            "POZO 1",
            df_eq["ubicacion_norm"].isin(["POSICION 04", "POSICION 05"])
        )
        agregar_subgrupo(
            "POZO 2",
            df_eq["ubicacion_norm"].isin(["POSICION 01", "POSICION 02", "POSICION 03"])
        )

    if not subgrupos:
        agregar_subgrupo("BOMBAS", df_eq["codigo"].astype(str).str.strip() != "")

    return subgrupos


def filtrar_matriz_subgrupo(df_matriz_nv, codigos):

    if df_matriz_nv.empty:
        return pd.DataFrame()

    codigos_set = set([str(c).strip() for c in codigos])

    return df_matriz_nv[
        df_matriz_nv["codigo"].astype(str).str.strip().isin(codigos_set)
    ].copy()


def crear_resumen_diario_subgrupo(df_matriz_subgrupo):

    if df_matriz_subgrupo.empty:
        return pd.DataFrame()

    registros = []

    for fecha, df_dia in df_matriz_subgrupo.groupby("fecha"):

        bombas_intervenidas = df_dia[
            (df_dia["cantidad_fallas_directas"] > 0) |
            (df_dia["horas_reales_intervencion"] > 0)
        ]["codigo"].astype(str).tolist()

        bombas_afectadas = df_dia[
            (df_dia["cantidad_fallas_afectan"] > 0) |
            (df_dia["horas_afectadas"] > 0)
        ]["codigo"].astype(str).tolist()

        horas_reales = float(df_dia["horas_reales_intervencion"].sum())
        horas_afectadas = float(df_dia["horas_afectadas"].sum())
        fallas = int(df_dia["cantidad_fallas_directas"].sum())
        df_dia_operativas = df_dia[df_dia.get("estado_operativo", pd.Series(dtype=str)).astype(str).str.upper() == "OPERATIVA"].copy()
        disponibilidad = float(df_dia_operativas["disponibilidad"].mean()) if not df_dia_operativas.empty else 100.0
        bombas_standby = df_dia[df_dia.get("estado_operativo", pd.Series(dtype=str)).astype(str).str.upper() == "STAND BY"]["codigo"].astype(str).tolist()

        registros.append({
            "fecha": fecha,
            "bombas_intervenidas": ", ".join(sorted(set(bombas_intervenidas))),
            "bombas_afectadas": ", ".join(sorted(set(bombas_afectadas))),
            "bombas_standby": ", ".join(sorted(set(bombas_standby))),
            "horas_reales": round(horas_reales, 2),
            "horas_afectadas": round(horas_afectadas, 2),
            "cantidad_fallas": fallas,
            "disponibilidad": round(disponibilidad, 2)
        })

    return pd.DataFrame(registros).sort_values("fecha")


def grafico_pdf_disponibilidad_subgrupo(df_matriz_subgrupo, nivel, nombre_subgrupo):

    if df_matriz_subgrupo.empty:
        return None

    df_plot = df_matriz_subgrupo.copy()
    fechas = sorted(df_plot["fecha"].unique().tolist())

    columnas_bombas = ["ubicacion", "codigo"]
    if "estado_operativo" in df_plot.columns:
        columnas_bombas.append("estado_operativo")

    df_bombas = (
        df_plot[columnas_bombas]
        .drop_duplicates()
        .sort_values(["ubicacion", "codigo"])
    )

    if df_bombas.empty:
        return None

    codigos = df_bombas["codigo"].astype(str).tolist()
    if "estado_operativo" in df_bombas.columns:
        etiquetas = [
            f"{str(row['ubicacion'])} (STAND BY)" if normalizar_texto(row.get("estado_operativo", "")) == "STAND BY" else str(row["ubicacion"])
            for _, row in df_bombas.iterrows()
        ]
    else:
        etiquetas = df_bombas["ubicacion"].astype(str).tolist()

    # Figura un poco más alta para dar aire a título, leyenda y etiquetas
    fig, ax1 = plt.subplots(figsize=(7.4, 4.15))

    x = list(range(len(fechas)))
    n = max(len(codigos), 1)

    # Se reduce ligeramente el ancho total del bloque para separar mejor las barras
    ancho = min(0.22, 0.68 / n)

    for i, codigo in enumerate(codigos):

        valores = []

        for fecha in fechas:
            df_val = df_plot[
                (df_plot["fecha"] == fecha) &
                (df_plot["codigo"].astype(str) == str(codigo))
            ]
            if df_val.empty:
                valores.append(100.0)
            else:
                valores.append(float(df_val.iloc[0]["disponibilidad"]))

        posiciones = [p + (i - (n - 1) / 2) * ancho for p in x]

        barras = ax1.bar(
            posiciones,
            valores,
            width=ancho * 0.92,
            edgecolor="black",
            linewidth=0.7,
            label=etiquetas[i]
        )

        # Etiquetas verticales dentro de la barra para evitar superposición
        for barra in barras:
            altura = barra.get_height()

            if altura <= 0:
                continue

            if altura >= 18:
                y_texto = max(8, altura - 8)
                va_texto = "top"
                color_texto = "white"
                bbox_texto = dict(
                    facecolor="black",
                    edgecolor="none",
                    alpha=0.55,
                    pad=1.1
                )
            else:
                y_texto = altura + 2
                va_texto = "bottom"
                color_texto = "black"
                bbox_texto = dict(
                    facecolor="white",
                    edgecolor="none",
                    alpha=0.80,
                    pad=1.0
                )

            ax1.text(
                barra.get_x() + barra.get_width() / 2,
                y_texto,
                f"{altura:.1f}%",
                ha="center",
                va=va_texto,
                rotation=90,
                fontsize=5.8,
                fontweight="bold",
                color=color_texto,
                bbox=bbox_texto,
                clip_on=True,
                zorder=9
            )

    ax1.set_title(
        f"Disponibilidad Mecánica {nivel} - {nombre_subgrupo}",
        fontsize=12,
        fontweight="bold",
        pad=8
    )
    ax1.set_xlabel("Fecha", fontsize=9)
    ax1.set_ylabel("% Disponibilidad mecánica", fontsize=9)
    ax1.set_ylim(0, 112)
    ax1.set_xticks(x)
    ax1.set_xticklabels([pd.to_datetime(f).strftime("%d/%m") for f in fechas], fontsize=8)
    ax1.grid(axis="y", linestyle="--", alpha=0.35)

    ax2 = ax1.twinx()

    fallas_por_fecha = []
    for fecha in fechas:
        fallas_por_fecha.append(
            int(df_plot[df_plot["fecha"] == fecha]["cantidad_fallas_directas"].sum())
        )

    ax2.plot(
        x,
        fallas_por_fecha,
        marker="o",
        markersize=5,
        linewidth=2,
        color="crimson",
        markerfacecolor="white",
        markeredgecolor="black",
        markeredgewidth=1,
        label="Cantidad de fallas reales",
        zorder=4
    )

    # Escala secundaria con más espacio para que la línea/etiquetas de fallas
    # no choquen con los porcentajes verticales de disponibilidad.
    max_fallas = max(fallas_por_fecha) if fallas_por_fecha else 0
    ax2.set_ylim(0, max(max_fallas + 2, 3))

    ax2.set_ylabel("Cantidad de fallas reales", fontsize=9)

    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()

    ax1.legend(
        handles1 + handles2,
        labels1 + labels2,
        loc="upper center",
        bbox_to_anchor=(0.5, -0.24),
        ncol=min(len(labels1) + len(labels2), 4),
        fontsize=6.6,
        frameon=True
    )

    fig.tight_layout(rect=[0, 0.08, 1, 0.98])

    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180)
    plt.close(fig)
    buffer.seek(0)

    return buffer


def grafico_pdf_paradas_subgrupo(df_matriz_subgrupo, nivel, nombre_subgrupo):

    if df_matriz_subgrupo.empty:
        return None

    df_plot = df_matriz_subgrupo.copy()
    fechas = sorted(df_plot["fecha"].unique().tolist())

    columnas_bombas = ["ubicacion", "codigo"]
    if "estado_operativo" in df_plot.columns:
        columnas_bombas.append("estado_operativo")

    df_bombas = (
        df_plot[columnas_bombas]
        .drop_duplicates()
        .sort_values(["ubicacion", "codigo"])
    )

    if df_bombas.empty:
        return None

    codigos = df_bombas["codigo"].astype(str).tolist()
    if "estado_operativo" in df_bombas.columns:
        etiquetas = [
            f"{str(row['ubicacion'])} (STAND BY)" if normalizar_texto(row.get("estado_operativo", "")) == "STAND BY" else str(row["ubicacion"])
            for _, row in df_bombas.iterrows()
        ]
    else:
        etiquetas = df_bombas["ubicacion"].astype(str).tolist()

    fig, ax = plt.subplots(figsize=(7.4, 4.15))

    x = list(range(len(fechas)))
    n = max(len(codigos), 1)
    ancho = min(0.22, 0.68 / n)
    max_horas = 0.0

    for i, codigo in enumerate(codigos):

        valores = []

        for fecha in fechas:
            df_val = df_plot[
                (df_plot["fecha"] == fecha) &
                (df_plot["codigo"].astype(str) == str(codigo))
            ]
            if df_val.empty:
                valores.append(0.0)
            else:
                valores.append(float(df_val.iloc[0]["horas_afectadas"]))

        if valores:
            max_horas = max(max_horas, max(valores))

        posiciones = [p + (i - (n - 1) / 2) * ancho for p in x]

        barras = ax.bar(
            posiciones,
            valores,
            width=ancho * 0.92,
            edgecolor="black",
            linewidth=0.7,
            label=etiquetas[i]
        )

        # Etiquetas verticales dentro de la barra para evitar que se monten entre bombas vecinas
        for barra in barras:
            altura = barra.get_height()

            if altura <= 0:
                continue

            if altura >= 1.0:
                y_texto = max(altura - max(max_horas * 0.06, 0.25), altura * 0.55)
                va_texto = "top"
                color_texto = "white"
                bbox_texto = dict(
                    facecolor="black",
                    edgecolor="none",
                    alpha=0.55,
                    pad=1.0
                )
            else:
                y_texto = altura + max(max_horas * 0.04, 0.05)
                va_texto = "bottom"
                color_texto = "black"
                bbox_texto = dict(
                    facecolor="white",
                    edgecolor="none",
                    alpha=0.80,
                    pad=1.0
                )

            ax.text(
                barra.get_x() + barra.get_width() / 2,
                y_texto,
                f"{altura:.1f} h",
                ha="center",
                va=va_texto,
                rotation=90,
                fontsize=5.8,
                fontweight="bold",
                color=color_texto,
                bbox=bbox_texto,
                clip_on=True
            )

    ax.set_title(
        f"Tiempo de Parada {nivel} - {nombre_subgrupo}",
        fontsize=12,
        fontweight="bold",
        pad=8
    )
    ax.set_xlabel("Fecha", fontsize=9)
    ax.set_ylabel("Horas de parada afectadas", fontsize=9)
    ax.set_xticks(x)
    ax.set_xticklabels([pd.to_datetime(f).strftime("%d/%m") for f in fechas], fontsize=8)
    ax.set_ylim(0, max(max_horas * 1.30, 1))
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.5, -0.24),
        ncol=min(len(etiquetas), 4),
        fontsize=6.6,
        frameon=True
    )

    fig.tight_layout(rect=[0, 0.08, 1, 0.98])

    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180)
    plt.close(fig)
    buffer.seek(0)

    return buffer


def grafico_pdf_disponibilidad_nivel(df_resumen_diario, nivel):

    if df_resumen_diario.empty:
        return None

    df_disp_dia = df_resumen_diario.copy()

    fig, ax1 = plt.subplots(figsize=(7.2, 3.9))

    x = list(range(len(df_disp_dia)))

    barras = ax1.bar(
        x,
        df_disp_dia["disponibilidad"],
        width=0.55,
        edgecolor="black",
        linewidth=0.8,
        label=f"Nivel {nivel}"
    )

    for barra in barras:
        altura = barra.get_height()
        ax1.text(
            barra.get_x() + barra.get_width() / 2,
            altura + 1.2,
            f"{altura:.1f}%",
            ha="center",
            va="bottom",
            fontsize=7,
            fontweight="bold"
        )

    ax1.set_title(f"Disponibilidad Mecánica del Nivel - {nivel}", fontsize=12, fontweight="bold")
    ax1.set_xlabel("Fecha", fontsize=9)
    ax1.set_ylabel("% Disponibilidad mecánica del nivel", fontsize=9)
    ax1.set_ylim(0, 112)
    ax1.set_xticks(x)
    ax1.set_xticklabels([pd.to_datetime(f).strftime("%d/%m") for f in df_disp_dia["fecha"]], fontsize=8)
    ax1.grid(axis="y", linestyle="--", alpha=0.35)

    ax2 = ax1.twinx()

    ax2.plot(
        x,
        df_disp_dia["cantidad_fallas"],
        marker="o",
        markersize=5,
        linewidth=2,
        color="crimson",
        markerfacecolor="white",
        markeredgecolor="black",
        markeredgewidth=1,
        label="Cantidad de fallas reales"
    )

    for x_pos, valor in enumerate(df_disp_dia["cantidad_fallas"]):
        ax2.annotate(
            str(int(valor)),
            (x_pos, valor),
            textcoords="offset points",
            xytext=(0, 8),
            ha="center",
            fontsize=7,
            fontweight="bold",
            bbox=dict(facecolor="white", edgecolor="black", linewidth=0.3, alpha=0.85, pad=1.2)
        )

    ax2.set_ylabel("Cantidad de fallas reales", fontsize=9)
    ax2.set_ylim(0, max(df_disp_dia["cantidad_fallas"].max() + 1, 2))

    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(handles1 + handles2, labels1 + labels2, loc="upper center", bbox_to_anchor=(0.5, -0.23), ncol=2, fontsize=7)

    fig.tight_layout(rect=[0, 0.05, 1, 1])

    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=170)
    plt.close(fig)
    buffer.seek(0)

    return buffer


def grafico_pdf_paradas_nivel(df_resumen_diario, nivel):

    if df_resumen_diario.empty:
        return None

    df_parada_dia = df_resumen_diario.copy()

    fig, ax = plt.subplots(figsize=(7.2, 3.9))

    x = list(range(len(df_parada_dia)))

    barras = ax.bar(
        x,
        df_parada_dia["horas_equipo"],
        width=0.55,
        edgecolor="black",
        linewidth=0.8,
        label=f"Nivel {nivel}"
    )

    max_horas = df_parada_dia["horas_equipo"].max() if not df_parada_dia.empty else 0

    for barra in barras:
        altura = barra.get_height()
        if altura > 0:
            ax.text(
                barra.get_x() + barra.get_width() / 2,
                altura + max(max_horas * 0.03, 0.05),
                f"{altura:.1f} h-eq",
                ha="center",
                va="bottom",
                fontsize=7,
                fontweight="bold"
            )

    ax.set_title(f"Tiempo de Parada Equivalente - {nivel}", fontsize=12, fontweight="bold")
    ax.set_xlabel("Fecha", fontsize=9)
    ax.set_ylabel("Horas-equipo de parada", fontsize=9)
    ax.set_xticks(x)
    ax.set_xticklabels([pd.to_datetime(f).strftime("%d/%m") for f in df_parada_dia["fecha"]], fontsize=8)
    ax.set_ylim(0, max(max_horas * 1.25, 1))
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.23), fontsize=7)

    fig.tight_layout(rect=[0, 0.05, 1, 1])

    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=170)
    plt.close(fig)
    buffer.seek(0)

    return buffer

def grafico_pdf_preventivo_correctivo(df_pdf, fecha_inicio, fecha_fin):

    if df_pdf.empty:
        return None

    df = df_pdf.copy()
    df["fecha_dia"] = df["fecha"].dt.date

    fechas = pd.date_range(fecha_inicio, fecha_fin, freq="D").date

    preventivo = []
    correctivo = []

    for fecha in fechas:
        df_fecha = df[df["fecha_dia"] == fecha]
        preventivo.append(df_fecha[df_fecha["tipo_mantenimiento"].str.upper() == "PREVENTIVO"]["tiempo_parada"].sum())
        correctivo.append(df_fecha[df_fecha["tipo_mantenimiento"].str.upper() == "CORRECTIVO"]["tiempo_parada"].sum())

    fig, ax1 = plt.subplots(figsize=(7.2, 3.9))

    x = list(range(len(fechas)))
    barras_prev = ax1.bar(x, preventivo, label="Horas preventivas", edgecolor="black", linewidth=0.8)

    max_prev = max(preventivo) if preventivo else 0
    max_corr = max(correctivo) if correctivo else 0

    for barra in barras_prev:
        altura = barra.get_height()
        if altura > 0:
            ax1.text(barra.get_x() + barra.get_width() / 2, altura + max(max_prev * 0.03, 0.05), f"{altura:.1f} h", ha="center", va="bottom", fontsize=7, fontweight="bold")

    ax1.set_title("Trabajos Preventivos vs Correctivos", fontsize=12, fontweight="bold")
    ax1.set_xlabel("Fecha", fontsize=9)
    ax1.set_ylabel("Horas preventivas", fontsize=9)
    ax1.set_xticks(x)
    ax1.set_xticklabels([pd.to_datetime(f).strftime("%d/%m") for f in fechas], fontsize=8)
    ax1.set_ylim(0, max(max_prev * 1.25, 1))
    ax1.grid(axis="y", linestyle="--", alpha=0.35)

    ax2 = ax1.twinx()
    ax2.plot(x, correctivo, marker="o", markersize=5, linewidth=2, color="crimson", markerfacecolor="white", markeredgecolor="black", markeredgewidth=1, label="Horas correctivas")

    for x_pos, valor in enumerate(correctivo):
        if valor > 0:
            ax2.annotate(f"{valor:.1f} h", (x_pos, valor), textcoords="offset points", xytext=(0, 8), ha="center", fontsize=7, fontweight="bold", bbox=dict(facecolor="white", edgecolor="black", linewidth=0.3, alpha=0.85, pad=1.2))

    ax2.set_ylabel("Horas correctivas", fontsize=9)
    ax2.set_ylim(0, max(max_corr * 1.25, 1))

    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(handles1 + handles2, labels1 + labels2, loc="upper center", bbox_to_anchor=(0.5, -0.23), ncol=2, fontsize=7)

    fig.tight_layout(rect=[0, 0.05, 1, 1])

    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=170)
    plt.close(fig)
    buffer.seek(0)

    return buffer


# Alias para compatibilidad con funciones antiguas del dashboard interactivo.
def grafico_pdf_disponibilidad(df_pdf, equipos_por_nivel, fecha_inicio, fecha_fin):
    df_equipos = preparar_equipos_detalle()
    niveles = sorted(df_pdf["nivel"].dropna().unique().tolist())
    if not niveles:
        return None
    df_resumen_diario = crear_resumen_diario_nivel(df_pdf, df_equipos, niveles[0], fecha_inicio, fecha_fin)
    return grafico_pdf_disponibilidad_nivel(df_resumen_diario, niveles[0])


def grafico_pdf_paradas(df_pdf, fecha_inicio, fecha_fin):
    df_equipos = preparar_equipos_detalle()
    niveles = sorted(df_pdf["nivel"].dropna().unique().tolist())
    if not niveles:
        return None
    df_resumen_diario = crear_resumen_diario_nivel(df_pdf, df_equipos, niveles[0], fecha_inicio, fecha_fin)
    return grafico_pdf_paradas_nivel(df_resumen_diario, niveles[0])


# ==========================================================
# PDF COMPLETO
# ==========================================================

def generar_pdf_informe_bombeo(df_informe, equipos_por_nivel, fecha_inicio, fecha_fin, incluir_fotos=False):

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.0 * cm,
        leftMargin=1.0 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm
    )

    styles = getSampleStyleSheet()

    titulo_style = ParagraphStyle("TituloVolcan", parent=styles["Title"], fontSize=16, textColor=colors.HexColor("#B71C1C"), spaceAfter=10)
    subtitulo_style = ParagraphStyle("SubtituloVolcan", parent=styles["Heading2"], fontSize=11, textColor=colors.HexColor("#212121"), spaceBefore=8, spaceAfter=6)
    normal_style = ParagraphStyle("NormalVolcan", parent=styles["Normal"], fontSize=8, leading=10)
    detalle_style = ParagraphStyle("Detalle", parent=styles["Normal"], fontSize=6, leading=7)

    elementos = []
    fecha_emision = datetime.now().strftime("%d/%m/%Y %H:%M")

    elementos.append(Paragraph("INFORME SEMANAL DEL SISTEMA DE BOMBEO", titulo_style))
    elementos.append(Paragraph(f"Periodo evaluado: {fecha_inicio.strftime('%d/%m/%Y')} al {fecha_fin.strftime('%d/%m/%Y')}", normal_style))
    elementos.append(Paragraph(f"Fecha de emisión: {fecha_emision}", normal_style))
    elementos.append(Spacer(1, 8))

    elementos.append(Paragraph(
        "Estimados,<br/><br/>Se adjunta el informe semanal del sistema de bombeo correspondiente al periodo seleccionado. "
        "El reporte consolida disponibilidad mecánica, cantidad de fallas, horas de parada y detalle de intervenciones. "
        "Para niveles con trenes, la parada de una bomba del tren se considera como afectación operativa del tren completo. "
        "Las bombas declaradas en STAND BY se muestran en el informe, pero no castigan la disponibilidad mecánica operativa.",
        normal_style
    ))
    elementos.append(Spacer(1, 8))

    df_equipos = preparar_equipos_detalle()

    if df_informe.empty:
        elementos.append(Paragraph("No se registran eventos para el periodo seleccionado.", normal_style))
        doc.build(elementos)
        pdf = buffer.getvalue()
        buffer.close()
        return pdf

    df_pdf = df_informe.copy()
    df_pdf["fecha_dia"] = df_pdf["fecha"].dt.date

    dias_periodo = (fecha_fin - fecha_inicio).days + 1
    total_eventos = len(df_pdf)
    total_horas = df_pdf["tiempo_parada"].sum()
    total_correctivo = df_pdf[df_pdf["tipo_mantenimiento"].str.upper() == "CORRECTIVO"]["tiempo_parada"].sum()
    total_preventivo = df_pdf[df_pdf["tipo_mantenimiento"].str.upper() == "PREVENTIVO"]["tiempo_parada"].sum()

    niveles = sorted(df_equipos["nivel"].dropna().unique().tolist())
    if not niveles:
        niveles = sorted(df_pdf["nivel"].dropna().unique().tolist())

    total_hp = 0
    total_horas_equivalentes = 0

    for nv in niveles:
        df_matriz_nv = crear_matriz_disponibilidad_nivel(df_pdf, df_equipos, nv, fecha_inicio, fecha_fin)
        if not df_matriz_nv.empty:
            df_matriz_nv_operativas = df_matriz_nv[df_matriz_nv["estado_operativo"].astype(str).str.upper() == "OPERATIVA"]
            total_hp += df_matriz_nv_operativas.shape[0] * 24
            df_resumen_diario_nv = crear_resumen_diario_nivel(df_pdf, df_equipos, nv, fecha_inicio, fecha_fin)
            total_horas_equivalentes += df_resumen_diario_nv["horas_equipo"].sum()

    dm_global = ((total_hp - total_horas_equivalentes) / total_hp) * 100 if total_hp > 0 else 0
    dm_global = max(0, min(dm_global, 100))

    total_bombas_instaladas = int(df_equipos["codigo"].nunique()) if not df_equipos.empty else 0
    total_bombas_standby = int(df_equipos[df_equipos["estado_operativo"].astype(str).str.upper() == "STAND BY"]["codigo"].nunique()) if not df_equipos.empty and "estado_operativo" in df_equipos.columns else 0
    total_bombas_operativas = max(total_bombas_instaladas - total_bombas_standby, 0)

    elementos.append(Paragraph("1. Resumen ejecutivo", subtitulo_style))
    resumen_data = [
        ["Eventos registrados", str(total_eventos)],
        ["Horas de parada registradas", f"{total_horas:.2f} h"],
        ["Horas de parada equivalentes", f"{total_horas_equivalentes:.2f} h"],
        ["Horas preventivas", f"{total_preventivo:.2f} h"],
        ["Horas correctivas", f"{total_correctivo:.2f} h"],
        ["Disponibilidad mecánica global operativa", f"{dm_global:.2f} %"],
        ["Bombas instaladas", str(total_bombas_instaladas)],
        ["Bombas operativas consideradas", str(total_bombas_operativas)],
        ["Bombas en stand by", str(total_bombas_standby)],
    ]
    elementos.append(crear_tabla_pdf(resumen_data, ["Indicador", "Valor"], font_size=7, col_widths=[9.0 * cm, 8.0 * cm]))
    elementos.append(Spacer(1, 8))

    elementos.append(Paragraph("2. Disponibilidad mecánica por nivel y grupo operativo", subtitulo_style))

    numero_grafica = 1

    for nv in niveles:

        df_eq_nv = df_equipos[
            df_equipos["nivel"].astype(str).str.strip() == str(nv).strip()
        ].copy()

        df_matriz_nv = crear_matriz_disponibilidad_nivel(
            df_pdf,
            df_equipos,
            nv,
            fecha_inicio,
            fecha_fin
        )

        if df_matriz_nv.empty or df_eq_nv.empty:
            continue

        subgrupos = obtener_subgrupos_nivel(df_eq_nv, nv)

        for subgrupo in subgrupos:

            df_matriz_sub = filtrar_matriz_subgrupo(
                df_matriz_nv,
                subgrupo["codigos"]
            )

            if df_matriz_sub.empty:
                continue

            df_resumen_sub = crear_resumen_diario_subgrupo(df_matriz_sub)
            grafico_dm = grafico_pdf_disponibilidad_subgrupo(
                df_matriz_sub,
                nv,
                subgrupo["nombre"]
            )

            tabla_diaria = []
            for _, row in df_resumen_sub.iterrows():
                tabla_diaria.append([
                    pd.to_datetime(row["fecha"]).strftime("%d/%m/%Y"),
                    str(row["bombas_intervenidas"]) if str(row["bombas_intervenidas"]).strip() else "N/I",
                    str(row["bombas_afectadas"]) if str(row["bombas_afectadas"]).strip() else "N/I",
                    f"{row['horas_reales']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(int(row["cantidad_fallas"])),
                    f"{row['disponibilidad']:.2f} %"
                ])

            df_detalle_bombas = crear_detalle_diario_bombas(df_matriz_sub)
            tabla_bombas = []
            for _, row in df_detalle_bombas.iterrows():
                tabla_bombas.append([
                    pd.to_datetime(row["fecha"]).strftime("%d/%m/%Y"),
                    str(row["ubicacion"]),
                    str(row["codigo"]),
                    str(row.get("estado_operativo", "OPERATIVA")),
                    f"{row['disponibilidad']:.2f} %",
                    str(int(row["cantidad_fallas_directas"])),
                    f"{row['horas_reales_intervencion']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(row.get("tipo_falla", "N/I")),
                    str(row.get("causa_preliminar", "N/I")),
                    str(row["intervencion"])
                ])

            bloque = [
                Paragraph(
                    f"2.{numero_grafica}. Nivel {nv} - {subgrupo['nombre']}",
                    subtitulo_style
                )
            ]

            if grafico_dm is not None:
                bloque.append(RLImage(grafico_dm, width=17.2 * cm, height=8.8 * cm))
                bloque.append(Spacer(1, 5))

            bloque.append(Paragraph("Resumen diario del grupo operativo", normal_style))
            bloque.append(crear_tabla_pdf(
                tabla_diaria,
                ["Fecha", "Bombas interv.", "Bombas afect.", "H. reales", "H. afect.", "Fallas", "% Disp."],
                font_size=4.9,
                col_widths=[1.55 * cm, 4.45 * cm, 4.65 * cm, 1.35 * cm, 1.15 * cm, 0.95 * cm, 1.35 * cm]
            ))
            bloque.append(Spacer(1, 5))

            bloque.append(Paragraph("Detalle diario por bomba del grupo", normal_style))
            bloque.append(crear_tabla_pdf(
                tabla_bombas,
                ["Fecha", "Ubicación", "Código", "Estado Op.", "% Disp.", "Fallas", "H. real", "H. afect.", "Tipo falla", "Causa", "Condición"],
                font_size=3.9,
                col_widths=[1.15 * cm, 1.65 * cm, 1.55 * cm, 1.05 * cm, 0.9 * cm, 0.65 * cm, 0.8 * cm, 0.85 * cm, 1.6 * cm, 1.65 * cm, 1.55 * cm]
            ))
            bloque.append(Spacer(1, 8))

            elementos.append(KeepTogether(bloque))
            numero_grafica += 1

    elementos.append(PageBreak())
    elementos.append(Paragraph("2.6. Resumen general de disponibilidad por nivel, grupo y bomba", subtitulo_style))
    elementos.append(Paragraph(
        "Para mejorar la lectura gerencial, el resumen se presenta separado por nivel. "
        "Las bombas en STAND BY se muestran como respaldo operativo, pero no se mezclan con el cálculo operativo del nivel.",
        normal_style
    ))
    elementos.append(Spacer(1, 6))

    df_resumen_general = crear_resumen_general_disponibilidad(
        df_pdf,
        df_equipos,
        niveles,
        fecha_inicio,
        fecha_fin
    )

    if not df_resumen_general.empty:

        for nv in niveles:

            df_resumen_nv = df_resumen_general[
                df_resumen_general["nivel"].astype(str).str.strip() == str(nv).strip()
            ].copy()

            if df_resumen_nv.empty:
                continue

            elementos.append(Paragraph(f"Resumen general - {nv}", normal_style))

            tabla_nivel = []
            for _, row in df_resumen_nv.iterrows():
                tabla_nivel.append([
                    str(row["tipo"]),
                    str(row["grupo"]),
                    str(row["ubicacion"]),
                    str(row["codigo"]),
                    texto_estado_operativo_resumen(row.get("estado_operativo", "")),
                    f"{row['horas_reales']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(int(row["fallas"])),
                    f"{row['disponibilidad']:.2f} %",
                    texto_condicion_resumen(row["condicion"])
                ])

            elementos.append(crear_tabla_pdf(
                tabla_nivel,
                ["Tipo", "Grupo", "Ubicación", "Código", "Estado Op.", "H. real", "H. afect.", "Fallas", "% Disp.", "Condición"],
                font_size=4.6,
                col_widths=[1.75 * cm, 1.75 * cm, 2.08 * cm, 1.45 * cm, 2.45 * cm, 0.75 * cm, 0.85 * cm, 0.55 * cm, 0.85 * cm, 2.65 * cm]
            ))
            elementos.append(Spacer(1, 8))

        df_resumen_sistema = df_resumen_general[
            df_resumen_general["tipo"].astype(str).str.upper() == "GLOBAL SISTEMA"
        ].copy()

        if not df_resumen_sistema.empty:
            elementos.append(Paragraph("Resumen global del sistema", normal_style))
            tabla_sistema = []
            for _, row in df_resumen_sistema.iterrows():
                tabla_sistema.append([
                    str(row["tipo"]),
                    str(row["grupo"]),
                    str(row["ubicacion"]),
                    str(row["codigo"]),
                    texto_estado_operativo_resumen(row.get("estado_operativo", "")),
                    f"{row['horas_reales']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(int(row["fallas"])),
                    f"{row['disponibilidad']:.2f} %",
                    texto_condicion_resumen(row["condicion"])
                ])

            elementos.append(crear_tabla_pdf(
                tabla_sistema,
                ["Tipo", "Grupo", "Ubicación", "Código", "Estado Op.", "H. real", "H. afect.", "Fallas", "% Disp.", "Condición"],
                font_size=4.6,
                col_widths=[1.75 * cm, 1.75 * cm, 2.08 * cm, 1.45 * cm, 2.45 * cm, 0.75 * cm, 0.85 * cm, 0.55 * cm, 0.85 * cm, 2.65 * cm]
            ))
    else:
        elementos.append(Paragraph("No se pudo generar el resumen general de disponibilidad.", normal_style))

    elementos.append(PageBreak())
    elementos.append(Paragraph("3. Tiempo de parada por grupo operativo", subtitulo_style))

    numero_grafica = 1

    for nv in niveles:

        df_eq_nv = df_equipos[
            df_equipos["nivel"].astype(str).str.strip() == str(nv).strip()
        ].copy()

        df_matriz_nv = crear_matriz_disponibilidad_nivel(
            df_pdf,
            df_equipos,
            nv,
            fecha_inicio,
            fecha_fin
        )

        if df_matriz_nv.empty or df_eq_nv.empty:
            continue

        subgrupos = obtener_subgrupos_nivel(df_eq_nv, nv)

        for subgrupo in subgrupos:

            df_matriz_sub = filtrar_matriz_subgrupo(
                df_matriz_nv,
                subgrupo["codigos"]
            )

            if df_matriz_sub.empty:
                continue

            df_resumen_sub = crear_resumen_diario_subgrupo(df_matriz_sub)
            grafico_paradas = grafico_pdf_paradas_subgrupo(
                df_matriz_sub,
                nv,
                subgrupo["nombre"]
            )

            tabla_data = []
            for _, row in df_resumen_sub.iterrows():
                tabla_data.append([
                    pd.to_datetime(row["fecha"]).strftime("%d/%m/%Y"),
                    str(row["bombas_intervenidas"]) if str(row["bombas_intervenidas"]).strip() else "N/I",
                    str(row["bombas_afectadas"]) if str(row["bombas_afectadas"]).strip() else "N/I",
                    f"{row['horas_reales']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(int(row["cantidad_fallas"]))
                ])

            bloque = [
                Paragraph(
                    f"3.{numero_grafica}. Nivel {nv} - {subgrupo['nombre']}",
                    subtitulo_style
                )
            ]

            if grafico_paradas is not None:
                bloque.append(RLImage(grafico_paradas, width=17.2 * cm, height=8.8 * cm))
                bloque.append(Spacer(1, 5))

            bloque.append(crear_tabla_pdf(
                tabla_data,
                ["Fecha", "Bombas interv.", "Bombas afect.", "H. reales", "H. afect.", "Fallas"],
                font_size=4.9,
                col_widths=[1.55 * cm, 4.65 * cm, 4.85 * cm, 1.35 * cm, 1.15 * cm, 0.95 * cm]
            ))
            bloque.append(Spacer(1, 8))

            elementos.append(KeepTogether(bloque))
            numero_grafica += 1

    elementos.append(PageBreak())

    grafico_pc = grafico_pdf_preventivo_correctivo(df_pdf, fecha_inicio, fecha_fin)
    if grafico_pc is not None:
        elementos.append(KeepTogether([
            Paragraph("4. Gráfico preventivo vs correctivo", subtitulo_style),
            RLImage(grafico_pc, width=17.2 * cm, height=8.8 * cm),
            Spacer(1, 6)
        ]))

    elementos.append(Paragraph("5. Trabajos preventivos vs correctivos", subtitulo_style))
    df_pc = df_pdf.pivot_table(index=["nivel", "fecha_dia"], columns="tipo_mantenimiento", values="tiempo_parada", aggfunc="sum", fill_value=0).reset_index()
    for col in ["PREVENTIVO", "CORRECTIVO"]:
        if col not in df_pc.columns:
            df_pc[col] = 0

    pc_data = []
    for _, row in df_pc.iterrows():
        pc_data.append([str(row["nivel"]), pd.to_datetime(row["fecha_dia"]).strftime("%d/%m/%Y"), f"{row['PREVENTIVO']:.2f}", f"{row['CORRECTIVO']:.2f}"])

    elementos.append(crear_tabla_pdf(pc_data, ["Nivel", "Fecha", "Horas preventivo", "Horas correctivo"], font_size=7))
    elementos.append(PageBreak())

    elementos.append(Paragraph("6. Análisis gerencial de intervenciones, fallas y causas", subtitulo_style))
    elementos.append(Paragraph(
        "Este resumen separa las actividades de mantenimiento ejecutadas de las fallas reales registradas durante el periodo evaluado, permitiendo diferenciar intervenciones preventivas/rutinarias de eventos correctivos que impactan la confiabilidad y disponibilidad del sistema de bombeo.",
        normal_style
    ))
    elementos.append(Spacer(1, 6))

    tipos_no_falla = [
        "INSPECCIÓN", "INSPECCION", "LIMPIEZA", "LUBRICACIÓN / ENGRASE",
        "LUBRICACIÓN", "LUBRICACION", "ENGRASE", "MANTENIMIENTO PREVENTIVO",
        "RUTINA PROGRAMADA", "CONDICIÓN NORMAL", "CONDICION NORMAL"
    ]

    if "tipo_falla" in df_pdf.columns:
        df_intervenciones = (
            df_pdf.groupby("tipo_falla", as_index=False)
            .agg(cantidad=("id", "count"), horas_parada=("tiempo_parada", "sum"))
            .sort_values(["cantidad", "horas_parada"], ascending=False)
        )
        data_intervenciones = [
            [str(row["tipo_falla"]), str(int(row["cantidad"])), f"{row['horas_parada']:.2f}"]
            for _, row in df_intervenciones.iterrows()
        ]
        elementos.append(Paragraph("6.1 Pareto de intervenciones registradas", normal_style))
        elementos.append(Paragraph(
            "Incluye inspecciones, limpieza, lubricación, mantenimiento preventivo y eventos correctivos registrados.",
            detalle_style
        ))
        elementos.append(crear_tabla_pdf(
            data_intervenciones,
            ["Intervención", "Cantidad", "Horas parada"],
            font_size=6.2,
            col_widths=[10.5 * cm, 2.4 * cm, 3.0 * cm]
        ))
        elementos.append(Spacer(1, 8))

        df_fallas_reales = df_pdf[
            ~df_pdf["tipo_falla"].astype(str).str.upper().str.strip().isin(tipos_no_falla)
        ].copy()

        if not df_fallas_reales.empty:
            df_pareto_falla_real = (
                df_fallas_reales.groupby("tipo_falla", as_index=False)
                .agg(cantidad=("id", "count"), horas_parada=("tiempo_parada", "sum"))
                .sort_values(["cantidad", "horas_parada"], ascending=False)
            )
            data_fallas_reales = [
                [str(row["tipo_falla"]), str(int(row["cantidad"])), f"{row['horas_parada']:.2f}"]
                for _, row in df_pareto_falla_real.iterrows()
            ]
            elementos.append(Paragraph("6.2 Pareto de fallas reales", normal_style))
            elementos.append(Paragraph(
                "Excluye inspecciones, limpieza, lubricación y actividades rutinarias; considera eventos correctivos que afectan la operación.",
                detalle_style
            ))
            elementos.append(crear_tabla_pdf(
                data_fallas_reales,
                ["Falla real", "Cantidad", "Horas parada"],
                font_size=6.2,
                col_widths=[10.5 * cm, 2.4 * cm, 3.0 * cm]
            ))
            elementos.append(Spacer(1, 8))
        else:
            elementos.append(Paragraph("No se registraron fallas reales en el periodo evaluado.", normal_style))
            elementos.append(Spacer(1, 8))

    if "causa_preliminar" in df_pdf.columns:
        df_pareto_causa = (
            df_pdf.groupby("causa_preliminar", as_index=False)
            .agg(cantidad=("id", "count"), horas_parada=("tiempo_parada", "sum"))
            .sort_values(["cantidad", "horas_parada"], ascending=False)
        )
        data_causas = [
            [str(row["causa_preliminar"]), str(int(row["cantidad"])), f"{row['horas_parada']:.2f}"]
            for _, row in df_pareto_causa.iterrows()
        ]
        elementos.append(Paragraph("6.3 Pareto de causas preliminares", normal_style))
        elementos.append(crear_tabla_pdf(
            data_causas,
            ["Causa preliminar", "Cantidad", "Horas parada"],
            font_size=6.2,
            col_widths=[10.5 * cm, 2.4 * cm, 3.0 * cm]
        ))

    elementos.append(PageBreak())

    elementos.append(Paragraph("7. Detalle de intervenciones", subtitulo_style))

    df_detalle = df_pdf[["fecha", "nivel", "ubicacion", "codigo", "hora_falla", "hora_subsanada", "descripcion"]].copy()
    df_detalle["fecha"] = df_detalle["fecha"].dt.strftime("%d/%m/%Y")

    detalle_data = []
    for _, row in df_detalle.iterrows():
        detalle_data.append([
            row["fecha"],
            str(row["nivel"]),
            str(row["ubicacion"]),
            str(row["codigo"]),
            str(row["hora_falla"]),
            str(row["hora_subsanada"]),
            Paragraph(str(row["descripcion"]), detalle_style)
        ])

    elementos.append(crear_tabla_pdf(
        detalle_data,
        ["FECHA", "NV", "UBICACIÓN", "CÓDIGO", "INICIO", "FIN", "DETALLE"],
        font_size=5.7,
        col_widths=[1.8 * cm, 1.1 * cm, 2.8 * cm, 2.4 * cm, 1.4 * cm, 1.4 * cm, 6.0 * cm]
    ))

    if incluir_fotos and "foto" in df_pdf.columns:
        df_fotos = df_pdf[df_pdf["foto"].astype(str).str.contains("base64", na=False)].copy()

        if not df_fotos.empty:
            elementos.append(PageBreak())
            elementos.append(Paragraph("8. Anexo fotográfico", subtitulo_style))
            elementos.append(Paragraph("Se adjuntan las evidencias fotográficas registradas para las intervenciones del periodo evaluado.", normal_style))
            elementos.append(Spacer(1, 8))

            for _, row in df_fotos.iterrows():
                imagen_bytes = obtener_bytes_imagen(row.get("foto", ""))
                if imagen_bytes is None:
                    continue

                try:
                    fecha_txt = pd.to_datetime(row.get("fecha")).strftime("%d/%m/%Y")
                except Exception:
                    fecha_txt = str(row.get("fecha", ""))

                tabla_evento = crear_tabla_pdf(
                    data=[[str(row.get("id", "")), fecha_txt, str(row.get("nivel", "")), str(row.get("codigo", "")), Paragraph(str(row.get("descripcion", "")), detalle_style)]],
                    headers=["EVENTO", "FECHA", "NV", "CÓDIGO", "DETALLE / TRABAJO"],
                    font_size=6,
                    col_widths=[2.1 * cm, 2.0 * cm, 1.4 * cm, 2.5 * cm, 8.8 * cm]
                )

                imagen_pdf = crear_imagen_pdf_desde_bytes(imagen_bytes, max_width=12.0 * cm, max_height=6.0 * cm)

                if imagen_pdf is not None:
                    elementos.append(KeepTogether([tabla_evento, Spacer(1, 6), imagen_pdf, Spacer(1, 12)]))

    doc.build(elementos)

    pdf = buffer.getvalue()
    buffer.close()
    return pdf


# ==========================================================
# WORD / DOCX
# ==========================================================

def set_cell_background(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=8, font_color="000000"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(font_color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def agregar_parrafo_word(documento, texto, font_size=9, bold=False, color="000000", align="left"):
    parrafo = documento.add_paragraph()
    if align == "center":
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        parrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = parrafo.add_run(str(texto))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    return parrafo


def agregar_titulo_word(documento, texto, nivel=1):
    parrafo = documento.add_paragraph()
    run = parrafo.add_run(str(texto))
    run.bold = True

    if nivel == 0:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(183, 28, 28)
    else:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

    return parrafo


def set_cell_width(cell, width_inches):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def agregar_tabla_word(documento, headers, data, font_size=7, column_widths=None, detalle_izquierda=True):

    tabla = documento.add_table(rows=1, cols=len(headers))
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"
    tabla.autofit = False
    tabla.allow_autofit = False

    header_cells = tabla.rows[0].cells
    for idx, header in enumerate(headers):
        if column_widths and idx < len(column_widths):
            set_cell_width(header_cells[idx], column_widths[idx])
            header_cells[idx].width = Inches(column_widths[idx])
        set_cell_background(header_cells[idx], "B71C1C")
        set_cell_text(header_cells[idx], header, bold=True, font_size=font_size, font_color="FFFFFF")

    for row_data in data:
        row = tabla.add_row()
        cells = row.cells

        texto_tipo = str(row_data[0]).strip().upper() if row_data else ""

        if texto_tipo == "GLOBAL GRUPO":
            color_fila = "D9EDF7"
            color_texto = "0C5460"
            negrita_fila = True
        elif texto_tipo == "GLOBAL NIVEL":
            color_fila = "B8DAFF"
            color_texto = "003366"
            negrita_fila = True
        elif texto_tipo == "GLOBAL SISTEMA":
            color_fila = "D4EDDA"
            color_texto = "155724"
            negrita_fila = True
        else:
            color_fila = "F5F5F5" if len(tabla.rows) % 2 == 0 else None
            color_texto = "000000"
            negrita_fila = False

        for idx, value in enumerate(row_data):
            cell = cells[idx]
            if column_widths and idx < len(column_widths):
                set_cell_width(cell, column_widths[idx])
                cell.width = Inches(column_widths[idx])

            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

            if detalle_izquierda and idx == len(row_data) - 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run(str(value))
            run.bold = negrita_fila
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(color_texto)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if color_fila:
                set_cell_background(cell, color_fila)

    documento.add_paragraph()
    return tabla


def agregar_imagen_word(documento, imagen_buffer, width_inches=6.3):
    try:
        imagen_buffer.seek(0)
        parrafo = documento.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = parrafo.add_run()
        run.add_picture(imagen_buffer, width=Inches(width_inches))
    except Exception:
        agregar_parrafo_word(documento, "No se pudo insertar la imagen en el informe Word.", font_size=8)


def generar_word_informe_bombeo(df_informe, equipos_por_nivel, fecha_inicio, fecha_fin, incluir_fotos=False):

    documento = Document()

    section = documento.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    fecha_emision = datetime.now().strftime("%d/%m/%Y %H:%M")

    agregar_titulo_word(documento, "INFORME SEMANAL DEL SISTEMA DE BOMBEO", nivel=0)
    agregar_parrafo_word(documento, f"Periodo evaluado: {fecha_inicio.strftime('%d/%m/%Y')} al {fecha_fin.strftime('%d/%m/%Y')}", font_size=8)
    agregar_parrafo_word(documento, f"Fecha de emisión: {fecha_emision}", font_size=8)
    agregar_parrafo_word(documento, "Estimados,", font_size=8)
    agregar_parrafo_word(
        documento,
        "Se adjunta el informe semanal del sistema de bombeo correspondiente al periodo seleccionado. El reporte consolida disponibilidad mecánica, cantidad de fallas, horas de parada y detalle de intervenciones. Para niveles con trenes, la parada de una bomba del tren se considera como afectación operativa del tren completo. Las bombas declaradas en STAND BY se muestran en el informe, pero no castigan la disponibilidad mecánica operativa.",
        font_size=8
    )

    if df_informe.empty:
        agregar_parrafo_word(documento, "No se registran eventos para el periodo seleccionado.", font_size=9, bold=True)
        buffer = BytesIO()
        documento.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    df_word = df_informe.copy()
    df_word["fecha_dia"] = df_word["fecha"].dt.date

    df_equipos = preparar_equipos_detalle()
    dias_periodo = (fecha_fin - fecha_inicio).days + 1
    total_eventos = len(df_word)
    total_horas = df_word["tiempo_parada"].sum()
    total_correctivo = df_word[df_word["tipo_mantenimiento"].str.upper() == "CORRECTIVO"]["tiempo_parada"].sum()
    total_preventivo = df_word[df_word["tipo_mantenimiento"].str.upper() == "PREVENTIVO"]["tiempo_parada"].sum()

    niveles = sorted(df_equipos["nivel"].dropna().unique().tolist())
    if not niveles:
        niveles = sorted(df_word["nivel"].dropna().unique().tolist())

    total_hp = 0
    total_horas_equivalentes = 0
    for nv in niveles:
        df_matriz_nv = crear_matriz_disponibilidad_nivel(df_word, df_equipos, nv, fecha_inicio, fecha_fin)
        if not df_matriz_nv.empty:
            df_matriz_nv_operativas = df_matriz_nv[df_matriz_nv["estado_operativo"].astype(str).str.upper() == "OPERATIVA"]
            total_hp += df_matriz_nv_operativas.shape[0] * 24
            df_resumen_diario_nv = crear_resumen_diario_nivel(df_word, df_equipos, nv, fecha_inicio, fecha_fin)
            total_horas_equivalentes += df_resumen_diario_nv["horas_equipo"].sum()

    dm_global = ((total_hp - total_horas_equivalentes) / total_hp) * 100 if total_hp > 0 else 0
    dm_global = max(0, min(dm_global, 100))

    total_bombas_instaladas = int(df_equipos["codigo"].nunique()) if not df_equipos.empty else 0
    total_bombas_standby = int(df_equipos[df_equipos["estado_operativo"].astype(str).str.upper() == "STAND BY"]["codigo"].nunique()) if not df_equipos.empty and "estado_operativo" in df_equipos.columns else 0
    total_bombas_operativas = max(total_bombas_instaladas - total_bombas_standby, 0)

    agregar_titulo_word(documento, "1. Resumen ejecutivo", nivel=1)
    agregar_tabla_word(
        documento,
        ["Indicador", "Valor"],
        [
            ["Eventos registrados", str(total_eventos)],
            ["Horas de parada registradas", f"{total_horas:.2f} h"],
            ["Horas de parada equivalentes", f"{total_horas_equivalentes:.2f} h"],
            ["Horas preventivas", f"{total_preventivo:.2f} h"],
            ["Horas correctivas", f"{total_correctivo:.2f} h"],
            ["Disponibilidad mecánica global operativa", f"{dm_global:.2f} %"],
            ["Bombas instaladas", str(total_bombas_instaladas)],
            ["Bombas operativas consideradas", str(total_bombas_operativas)],
            ["Bombas en stand by", str(total_bombas_standby)],
        ],
        font_size=7,
        column_widths=[3.8, 2.8]
    )

    agregar_titulo_word(documento, "2. Disponibilidad mecánica por nivel y grupo operativo", nivel=1)

    numero_grafica = 1

    for nv in niveles:

        df_eq_nv = df_equipos[
            df_equipos["nivel"].astype(str).str.strip() == str(nv).strip()
        ].copy()

        df_matriz_nv = crear_matriz_disponibilidad_nivel(
            df_word,
            df_equipos,
            nv,
            fecha_inicio,
            fecha_fin
        )

        if df_matriz_nv.empty or df_eq_nv.empty:
            continue

        subgrupos = obtener_subgrupos_nivel(df_eq_nv, nv)

        for subgrupo in subgrupos:

            df_matriz_sub = filtrar_matriz_subgrupo(
                df_matriz_nv,
                subgrupo["codigos"]
            )

            if df_matriz_sub.empty:
                continue

            df_resumen_sub = crear_resumen_diario_subgrupo(df_matriz_sub)

            agregar_titulo_word(
                documento,
                f"2.{numero_grafica}. Nivel {nv} - {subgrupo['nombre']}",
                nivel=1
            )

            grafico_dm = grafico_pdf_disponibilidad_subgrupo(
                df_matriz_sub,
                nv,
                subgrupo["nombre"]
            )

            if grafico_dm is not None:
                agregar_imagen_word(documento, grafico_dm, width_inches=6.5)

            tabla_diaria = []
            for _, row in df_resumen_sub.iterrows():
                tabla_diaria.append([
                    pd.to_datetime(row["fecha"]).strftime("%d/%m/%Y"),
                    str(row["bombas_intervenidas"]) if str(row["bombas_intervenidas"]).strip() else "N/I",
                    str(row["bombas_afectadas"]) if str(row["bombas_afectadas"]).strip() else "N/I",
                    f"{row['horas_reales']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(int(row["cantidad_fallas"])),
                    f"{row['disponibilidad']:.2f} %"
                ])

            agregar_parrafo_word(documento, "Resumen diario del grupo operativo", font_size=7, bold=True)
            agregar_tabla_word(
                documento,
                ["Fecha", "Bombas interv.", "Bombas afect.", "H. reales", "H. afect.", "Fallas", "% Disp."],
                tabla_diaria,
                font_size=4.8,
                column_widths=[0.62, 1.75, 1.85, 0.55, 0.48, 0.38, 0.52],
                detalle_izquierda=False
            )

            df_detalle_bombas = crear_detalle_diario_bombas(df_matriz_sub)
            tabla_bombas = []
            for _, row in df_detalle_bombas.iterrows():
                tabla_bombas.append([
                    pd.to_datetime(row["fecha"]).strftime("%d/%m/%Y"),
                    str(row["ubicacion"]),
                    str(row["codigo"]),
                    str(row.get("estado_operativo", "OPERATIVA")),
                    f"{row['disponibilidad']:.2f} %",
                    str(int(row["cantidad_fallas_directas"])),
                    f"{row['horas_reales_intervencion']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(row.get("tipo_falla", "N/I")),
                    str(row.get("causa_preliminar", "N/I")),
                    str(row["intervencion"])
                ])

            agregar_parrafo_word(documento, "Detalle diario por bomba del grupo", font_size=7, bold=True)
            agregar_tabla_word(
                documento,
                ["Fecha", "Ubicación", "Código", "Estado Op.", "% Disp.", "Fallas", "H. real", "H. afect.", "Tipo falla", "Causa", "Condición"],
                tabla_bombas,
                font_size=3.8,
                column_widths=[0.48, 0.72, 0.68, 0.55, 0.42, 0.28, 0.34, 0.36, 0.65, 0.72, 0.65],
                detalle_izquierda=False
            )

            numero_grafica += 1

    documento.add_page_break()
    agregar_titulo_word(documento, "2.6. Resumen general de disponibilidad por nivel, grupo y bomba", nivel=1)
    agregar_parrafo_word(
        documento,
        "Para mejorar la lectura gerencial, el resumen se presenta separado por nivel. Las bombas en STAND BY se muestran como respaldo operativo, pero no se mezclan con el cálculo operativo del nivel.",
        font_size=8
    )

    df_resumen_general = crear_resumen_general_disponibilidad(
        df_word,
        df_equipos,
        niveles,
        fecha_inicio,
        fecha_fin
    )

    if not df_resumen_general.empty:

        for nv in niveles:

            df_resumen_nv = df_resumen_general[
                df_resumen_general["nivel"].astype(str).str.strip() == str(nv).strip()
            ].copy()

            if df_resumen_nv.empty:
                continue

            agregar_parrafo_word(documento, f"Resumen general - {nv}", font_size=7, bold=True)

            tabla_nivel = []
            for _, row in df_resumen_nv.iterrows():
                tabla_nivel.append([
                    str(row["tipo"]),
                    str(row["grupo"]),
                    str(row["ubicacion"]),
                    str(row["codigo"]),
                    texto_estado_operativo_resumen(row.get("estado_operativo", "")),
                    f"{row['horas_reales']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(int(row["fallas"])),
                    f"{row['disponibilidad']:.2f} %",
                    texto_condicion_resumen(row["condicion"])
                ])

            agregar_tabla_word(
                documento,
                ["Tipo", "Grupo", "Ubicación", "Código", "Estado Op.", "H. real", "H. afect.", "Fallas", "% Disp.", "Condición"],
                tabla_nivel,
                font_size=3.7,
                column_widths=[0.46, 0.68, 0.78, 0.58, 1.15, 0.30, 0.34, 0.25, 0.36, 1.25],
                detalle_izquierda=False
            )

        df_resumen_sistema = df_resumen_general[
            df_resumen_general["tipo"].astype(str).str.upper() == "GLOBAL SISTEMA"
        ].copy()

        if not df_resumen_sistema.empty:
            agregar_parrafo_word(documento, "Resumen global del sistema", font_size=7, bold=True)

            tabla_sistema = []
            for _, row in df_resumen_sistema.iterrows():
                tabla_sistema.append([
                    str(row["tipo"]),
                    str(row["grupo"]),
                    str(row["ubicacion"]),
                    str(row["codigo"]),
                    texto_estado_operativo_resumen(row.get("estado_operativo", "")),
                    f"{row['horas_reales']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(int(row["fallas"])),
                    f"{row['disponibilidad']:.2f} %",
                    texto_condicion_resumen(row["condicion"])
                ])

            agregar_tabla_word(
                documento,
                ["Tipo", "Grupo", "Ubicación", "Código", "Estado Op.", "H. real", "H. afect.", "Fallas", "% Disp.", "Condición"],
                tabla_sistema,
                font_size=3.8,
                column_widths=[0.46, 0.68, 0.78, 0.58, 1.15, 0.30, 0.34, 0.25, 0.36, 1.25],
                detalle_izquierda=False
            )
    else:
        agregar_parrafo_word(documento, "No se pudo generar el resumen general de disponibilidad.", font_size=8, bold=True)

    documento.add_page_break()
    agregar_titulo_word(documento, "3. Tiempo de parada por grupo operativo", nivel=1)

    numero_grafica = 1

    for nv in niveles:

        df_eq_nv = df_equipos[
            df_equipos["nivel"].astype(str).str.strip() == str(nv).strip()
        ].copy()

        df_matriz_nv = crear_matriz_disponibilidad_nivel(
            df_word,
            df_equipos,
            nv,
            fecha_inicio,
            fecha_fin
        )

        if df_matriz_nv.empty or df_eq_nv.empty:
            continue

        subgrupos = obtener_subgrupos_nivel(df_eq_nv, nv)

        for subgrupo in subgrupos:

            df_matriz_sub = filtrar_matriz_subgrupo(
                df_matriz_nv,
                subgrupo["codigos"]
            )

            if df_matriz_sub.empty:
                continue

            df_resumen_sub = crear_resumen_diario_subgrupo(df_matriz_sub)

            agregar_titulo_word(
                documento,
                f"3.{numero_grafica}. Nivel {nv} - {subgrupo['nombre']}",
                nivel=1
            )

            grafico_paradas = grafico_pdf_paradas_subgrupo(
                df_matriz_sub,
                nv,
                subgrupo["nombre"]
            )

            if grafico_paradas is not None:
                agregar_imagen_word(documento, grafico_paradas, width_inches=6.5)

            tabla_data = []
            for _, row in df_resumen_sub.iterrows():
                tabla_data.append([
                    pd.to_datetime(row["fecha"]).strftime("%d/%m/%Y"),
                    str(row["bombas_intervenidas"]) if str(row["bombas_intervenidas"]).strip() else "N/I",
                    str(row["bombas_afectadas"]) if str(row["bombas_afectadas"]).strip() else "N/I",
                    f"{row['horas_reales']:.2f}",
                    f"{row['horas_afectadas']:.2f}",
                    str(int(row["cantidad_fallas"]))
                ])

            agregar_tabla_word(
                documento,
                ["Fecha", "Bombas interv.", "Bombas afect.", "H. reales", "H. afect.", "Fallas"],
                tabla_data,
                font_size=4.8,
                column_widths=[0.62, 1.9, 2.0, 0.55, 0.48, 0.38],
                detalle_izquierda=False
            )

            numero_grafica += 1

    documento.add_page_break()

    grafico_pc = grafico_pdf_preventivo_correctivo(df_word, fecha_inicio, fecha_fin)
    if grafico_pc is not None:
        agregar_titulo_word(documento, "4. Gráfico preventivo vs correctivo", nivel=1)
        agregar_imagen_word(documento, grafico_pc, width_inches=6.5)

    agregar_titulo_word(documento, "5. Trabajos preventivos vs correctivos", nivel=1)
    df_pc = df_word.pivot_table(index=["nivel", "fecha_dia"], columns="tipo_mantenimiento", values="tiempo_parada", aggfunc="sum", fill_value=0).reset_index()
    for col in ["PREVENTIVO", "CORRECTIVO"]:
        if col not in df_pc.columns:
            df_pc[col] = 0

    pc_data = []
    for _, row in df_pc.iterrows():
        pc_data.append([str(row["nivel"]), pd.to_datetime(row["fecha_dia"]).strftime("%d/%m/%Y"), f"{row['PREVENTIVO']:.2f}", f"{row['CORRECTIVO']:.2f}"])

    agregar_tabla_word(documento, ["Nivel", "Fecha", "Horas preventivo", "Horas correctivo"], pc_data, font_size=7, column_widths=[1.25, 1.45, 1.8, 1.8])

    agregar_titulo_word(documento, "6. Análisis gerencial de intervenciones, fallas y causas", nivel=1)
    agregar_parrafo_word(
        documento,
        "Este resumen separa las actividades de mantenimiento ejecutadas de las fallas reales registradas durante el periodo evaluado, permitiendo diferenciar intervenciones preventivas/rutinarias de eventos correctivos que impactan la confiabilidad y disponibilidad del sistema de bombeo.",
        font_size=8
    )

    tipos_no_falla = [
        "INSPECCIÓN", "INSPECCION", "LIMPIEZA", "LUBRICACIÓN / ENGRASE",
        "LUBRICACIÓN", "LUBRICACION", "ENGRASE", "MANTENIMIENTO PREVENTIVO",
        "RUTINA PROGRAMADA", "CONDICIÓN NORMAL", "CONDICION NORMAL"
    ]

    if "tipo_falla" in df_word.columns:
        df_intervenciones = (
            df_word.groupby("tipo_falla", as_index=False)
            .agg(cantidad=("id", "count"), horas_parada=("tiempo_parada", "sum"))
            .sort_values(["cantidad", "horas_parada"], ascending=False)
        )
        data_intervenciones = [
            [str(row["tipo_falla"]), str(int(row["cantidad"])), f"{row['horas_parada']:.2f}"]
            for _, row in df_intervenciones.iterrows()
        ]
        agregar_parrafo_word(documento, "6.1 Pareto de intervenciones registradas", font_size=7, bold=True)
        agregar_tabla_word(
            documento,
            ["Intervención", "Cantidad", "Horas parada"],
            data_intervenciones,
            font_size=6.2,
            column_widths=[4.6, 1.0, 1.2],
            detalle_izquierda=True
        )

        df_fallas_reales = df_word[
            ~df_word["tipo_falla"].astype(str).str.upper().str.strip().isin(tipos_no_falla)
        ].copy()

        if not df_fallas_reales.empty:
            df_pareto_falla_real = (
                df_fallas_reales.groupby("tipo_falla", as_index=False)
                .agg(cantidad=("id", "count"), horas_parada=("tiempo_parada", "sum"))
                .sort_values(["cantidad", "horas_parada"], ascending=False)
            )
            data_fallas_reales = [
                [str(row["tipo_falla"]), str(int(row["cantidad"])), f"{row['horas_parada']:.2f}"]
                for _, row in df_pareto_falla_real.iterrows()
            ]
            agregar_parrafo_word(documento, "6.2 Pareto de fallas reales", font_size=7, bold=True)
            agregar_tabla_word(
                documento,
                ["Falla real", "Cantidad", "Horas parada"],
                data_fallas_reales,
                font_size=6.2,
                column_widths=[4.6, 1.0, 1.2],
                detalle_izquierda=True
            )
        else:
            agregar_parrafo_word(documento, "No se registraron fallas reales en el periodo evaluado.", font_size=7, bold=True)

    if "causa_preliminar" in df_word.columns:
        df_pareto_causa = (
            df_word.groupby("causa_preliminar", as_index=False)
            .agg(cantidad=("id", "count"), horas_parada=("tiempo_parada", "sum"))
            .sort_values(["cantidad", "horas_parada"], ascending=False)
        )
        data_causas = [
            [str(row["causa_preliminar"]), str(int(row["cantidad"])), f"{row['horas_parada']:.2f}"]
            for _, row in df_pareto_causa.iterrows()
        ]
        agregar_parrafo_word(documento, "6.3 Pareto de causas preliminares", font_size=7, bold=True)
        agregar_tabla_word(
            documento,
            ["Causa preliminar", "Cantidad", "Horas parada"],
            data_causas,
            font_size=6.2,
            column_widths=[4.6, 1.0, 1.2],
            detalle_izquierda=True
        )

    agregar_titulo_word(documento, "7. Detalle de intervenciones", nivel=1)
    df_detalle = df_word[["fecha", "nivel", "ubicacion", "codigo", "hora_falla", "hora_subsanada", "descripcion"]].copy()
    df_detalle["fecha"] = df_detalle["fecha"].dt.strftime("%d/%m/%Y")

    detalle_data = []
    for _, row in df_detalle.iterrows():
        detalle_data.append([row["fecha"], str(row["nivel"]), str(row["ubicacion"]), str(row["codigo"]), str(row["hora_falla"]), str(row["hora_subsanada"]), str(row["descripcion"])])

    agregar_tabla_word(
        documento,
        ["FECHA", "NV", "UBICACIÓN", "CÓDIGO", "INICIO", "FIN", "DETALLE"],
        detalle_data,
        font_size=5.5,
        column_widths=[0.8, 0.55, 1.15, 1.0, 0.55, 0.55, 2.15],
        detalle_izquierda=True
    )

    if incluir_fotos and "foto" in df_word.columns:
        df_fotos = df_word[df_word["foto"].astype(str).str.contains("base64", na=False)].copy()
        if not df_fotos.empty:
            documento.add_page_break()
            agregar_titulo_word(documento, "8. Anexo fotográfico", nivel=1)
            agregar_parrafo_word(documento, "Se adjuntan las evidencias fotográficas registradas para las intervenciones del periodo evaluado.", font_size=8)

            for _, row in df_fotos.iterrows():
                imagen_bytes = obtener_bytes_imagen(row.get("foto", ""))
                if imagen_bytes is None:
                    continue

                try:
                    fecha_txt = pd.to_datetime(row.get("fecha")).strftime("%d/%m/%Y")
                except Exception:
                    fecha_txt = str(row.get("fecha", ""))

                agregar_tabla_word(
                    documento,
                    ["EVENTO", "FECHA", "NV", "CÓDIGO", "DETALLE / TRABAJO"],
                    [[str(row.get("id", "")), fecha_txt, str(row.get("nivel", "")), str(row.get("codigo", "")), str(row.get("descripcion", ""))]],
                    font_size=6,
                    column_widths=[0.9, 0.8, 0.55, 1.0, 3.0],
                    detalle_izquierda=True
                )

                imagen_buffer = BytesIO(imagen_bytes)
                agregar_imagen_word(documento, imagen_buffer, width_inches=3.7)

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================================
# DASHBOARD STREAMLIT
# ==========================================================

def mostrar_dashboard():

    st.title("📊 Dashboard de Eventos - Sistema Bombeo")
    st.caption("Mantenimiento Mecánico · Planeamiento · Confiabilidad")
    st.markdown("---")

    df = cargar_bitacora()

    if df.empty:
        st.warning("No existen registros aún.")
        return

    df.columns = df.columns.str.strip().str.lower()

    columnas_necesarias = [
        "id", "fecha", "tecnico", "apoyo_1", "apoyo_2", "sistema", "tipo_mantenimiento", "nivel", "ubicacion", "codigo", "hora_falla", "hora_subsanada", "tiempo_parada", "tipo_falla", "causa_preliminar", "repuesto_requerido", "descripcion", "estado", "foto"
    ]

    for col in columnas_necesarias:
        if col not in df.columns:
            df[col] = ""

    df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")
    df["tiempo_parada"] = pd.to_numeric(df["tiempo_parada"], errors="coerce").fillna(0)

    for col in ["codigo", "nivel", "ubicacion", "tipo_mantenimiento", "tipo_falla", "causa_preliminar", "repuesto_requerido", "estado", "tecnico", "apoyo_1", "apoyo_2"]:
        df[col] = df[col].astype(str).str.strip()

    df = df.dropna(subset=["fecha"])

    if df.empty:
        st.warning("No hay fechas válidas en la bitácora.")
        return

    st.subheader("🔎 Filtros")

    colf1, colf2, colf3, colf4 = st.columns(4)

    with colf1:
        nivel = st.selectbox("Nivel", ["TODOS"] + sorted(df["nivel"].dropna().unique().tolist()))
    with colf2:
        bomba = st.selectbox("Bomba", ["TODOS"] + sorted(df["codigo"].dropna().unique().tolist()))
    with colf3:
        falla = st.selectbox("Tipo de falla", ["TODOS"] + sorted(df["tipo_falla"].dropna().unique().tolist()))
    with colf4:
        causa = st.selectbox("Causa preliminar", ["TODOS"] + sorted(df["causa_preliminar"].dropna().unique().tolist()))

    colf5, colf6, colf7, colf8 = st.columns(4)

    with colf5:
        estado = st.selectbox("Estado", ["TODOS"] + sorted(df["estado"].dropna().unique().tolist()))
    with colf6:
        tipo_mantenimiento = st.selectbox("Tipo mantenimiento", ["TODOS"] + sorted(df["tipo_mantenimiento"].dropna().unique().tolist()))

    fecha_min = df["fecha"].min().date()
    fecha_max = df["fecha"].max().date()

    with colf7:
        fecha_inicio = st.date_input("Fecha inicio", value=fecha_min)
    with colf8:
        fecha_fin = st.date_input("Fecha fin", value=fecha_max)

    equipos_por_nivel = preparar_equipos_por_nivel()

    df_informe = df[(df["fecha"].dt.date >= fecha_inicio) & (df["fecha"].dt.date <= fecha_fin)].copy()

    st.markdown("---")
    st.header("📄 Informe Semanal - Sistema de Bombeo")
    st.caption("Exporta el informe formal en PDF y Word en formato vertical, con gráficas separadas por nivel.")

    incluir_fotos_pdf = st.checkbox(
        "Incluir anexo fotográfico en el PDF/Word",
        value=False,
        help="Activa esta opción solo cuando necesites sustentar el informe con evidencias. El archivo será más pesado."
    )

    pdf_bytes = generar_pdf_informe_bombeo(
        df_informe=df_informe,
        equipos_por_nivel=equipos_por_nivel,
        fecha_inicio=fecha_inicio,
        fecha_fin=fecha_fin,
        incluir_fotos=incluir_fotos_pdf
    )

    st.download_button(
        label="📄 Exportar informe semanal PDF",
        data=pdf_bytes,
        file_name=f"informe_bombeo_{fecha_inicio}_{fecha_fin}.pdf",
        mime="application/pdf",
        use_container_width=True
    )

    word_bytes = generar_word_informe_bombeo(
        df_informe=df_informe,
        equipos_por_nivel=equipos_por_nivel,
        fecha_inicio=fecha_inicio,
        fecha_fin=fecha_fin,
        incluir_fotos=incluir_fotos_pdf
    )

    st.download_button(
        label="📝 Exportar informe Word",
        data=word_bytes,
        file_name=f"informe_bombeo_{fecha_inicio}_{fecha_fin}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

    st.markdown("---")

    df_filtrado = df.copy()
    df_filtrado = df_filtrado[(df_filtrado["fecha"].dt.date >= fecha_inicio) & (df_filtrado["fecha"].dt.date <= fecha_fin)]

    if nivel != "TODOS":
        df_filtrado = df_filtrado[df_filtrado["nivel"] == nivel]
    if bomba != "TODOS":
        df_filtrado = df_filtrado[df_filtrado["codigo"] == bomba]
    if falla != "TODOS":
        df_filtrado = df_filtrado[df_filtrado["tipo_falla"] == falla]
    if causa != "TODOS":
        df_filtrado = df_filtrado[df_filtrado["causa_preliminar"] == causa]
    if estado != "TODOS":
        df_filtrado = df_filtrado[df_filtrado["estado"] == estado]
    if tipo_mantenimiento != "TODOS":
        df_filtrado = df_filtrado[df_filtrado["tipo_mantenimiento"] == tipo_mantenimiento]

    if df_filtrado.empty:
        st.warning("No hay registros con los filtros seleccionados.")
        return

    total_eventos = len(df_filtrado)
    total_horas = df_filtrado["tiempo_parada"].sum()
    mttr = df_filtrado["tiempo_parada"].mean()

    dias_periodo = (fecha_fin - fecha_inicio).days + 1
    bombas_involucradas = df_filtrado["codigo"].nunique()
    horas_calendario = dias_periodo * 24 * max(bombas_involucradas, 1)
    disponibilidad = ((horas_calendario - total_horas) / horas_calendario) * 100 if horas_calendario > 0 else 0
    disponibilidad = max(0, min(disponibilidad, 100))

    mtbf = round((horas_calendario - total_horas) / total_eventos, 2) if total_eventos > 1 else 0

    bomba_top = df_filtrado.groupby("codigo")["tiempo_parada"].sum().sort_values(ascending=False).index[0] if not df_filtrado.empty else "-"
    falla_top = df_filtrado["tipo_falla"].mode()[0] if not df_filtrado["tipo_falla"].empty else "-"
    causa_top = df_filtrado["causa_preliminar"].mode()[0] if not df_filtrado["causa_preliminar"].empty else "-"

    eventos_pendientes = df_filtrado[df_filtrado["estado"].str.upper().isin(["PENDIENTE", "FUERA DE SERVICIO", "EN SEGUIMIENTO"])].shape[0]
    repuestos_requeridos = df_filtrado[df_filtrado["repuesto_requerido"].astype(str).str.strip() != ""].shape[0]
    tecnicos_involucrados = df_filtrado["tecnico"].nunique()

    st.subheader("📌 Indicadores ejecutivos")

    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Eventos", total_eventos)
    k2.metric("Horas parada", round(total_horas, 2))
    k3.metric("Disponibilidad mecánica", f"{round(disponibilidad, 2)} %")
    k4.metric("MTTR promedio", f"{round(mttr, 2)} h")

    k5, k6, k7, k8 = st.columns(4)
    k5.metric("MTBF estimado", f"{mtbf} h")
    k6.metric("Bomba crítica", bomba_top)
    k7.metric("Eventos pendientes", eventos_pendientes)
    k8.metric("Repuestos requeridos", repuestos_requeridos)

    cinfo1, cinfo2, cinfo3 = st.columns(3)
    cinfo1.info(f"Falla más común: **{falla_top}**")
    cinfo2.info(f"Causa más común: **{causa_top}**")
    cinfo3.info(f"Técnicos involucrados: **{tecnicos_involucrados}**")

    st.markdown("---")
    st.subheader("📈 Análisis gráfico gerencial")

    st.markdown("### 1. Estado de eventos")
    df_estado = df_filtrado.groupby("estado").size().reset_index(name="cantidad").sort_values("cantidad", ascending=False)
    if not df_estado.empty:
        chart_estado = alt.Chart(df_estado).mark_arc(innerRadius=60).encode(theta=alt.Theta("cantidad:Q"), color=alt.Color("estado:N", title="Estado"), tooltip=["estado", "cantidad"])
        st.altair_chart(chart_estado, use_container_width=True)

    st.markdown("### 2. Top 10 bombas críticas por horas de parada")
    df_top_bombas = df_filtrado.groupby("codigo")["tiempo_parada"].sum().reset_index().sort_values("tiempo_parada", ascending=False).head(10)
    chart_top_bombas = alt.Chart(df_top_bombas).mark_bar().encode(x=alt.X("tiempo_parada:Q", title="Horas de parada"), y=alt.Y("codigo:N", sort="-x", title="Bomba"), color=alt.Color("codigo:N", legend=None), tooltip=["codigo", alt.Tooltip("tiempo_parada:Q", format=".2f")])
    st.altair_chart(chart_top_bombas, use_container_width=True)

    st.markdown("### 3. Pareto de fallas")
    df_pareto_falla = df_filtrado.groupby("tipo_falla").size().reset_index(name="cantidad").sort_values("cantidad", ascending=False)
    chart_pareto_falla = alt.Chart(df_pareto_falla).mark_bar().encode(x=alt.X("tipo_falla:N", sort="-y", title="Tipo de falla"), y=alt.Y("cantidad:Q", title="Cantidad"), color=alt.Color("tipo_falla:N", legend=None), tooltip=["tipo_falla", "cantidad"])
    st.altair_chart(chart_pareto_falla, use_container_width=True)

    st.markdown("### 4. Pareto de causas preliminares")
    df_pareto_causa = df_filtrado.groupby("causa_preliminar").size().reset_index(name="cantidad").sort_values("cantidad", ascending=False)
    chart_pareto_causa = alt.Chart(df_pareto_causa).mark_bar().encode(x=alt.X("causa_preliminar:N", sort="-y", title="Causa preliminar"), y=alt.Y("cantidad:Q", title="Cantidad"), color=alt.Color("causa_preliminar:N", legend=None), tooltip=["causa_preliminar", "cantidad"])
    st.altair_chart(chart_pareto_causa, use_container_width=True)

    st.markdown("### 5. Repuestos más solicitados")
    df_rep_chart = df_filtrado[df_filtrado["repuesto_requerido"].astype(str).str.strip() != ""].copy()
    if df_rep_chart.empty:
        st.info("No hay repuestos registrados para graficar.")
    else:
        df_rep_chart = df_rep_chart.groupby("repuesto_requerido").size().reset_index(name="cantidad").sort_values("cantidad", ascending=False)
        chart_repuestos = alt.Chart(df_rep_chart).mark_bar().encode(x=alt.X("cantidad:Q", title="Cantidad"), y=alt.Y("repuesto_requerido:N", sort="-x", title="Repuesto"), color=alt.Color("repuesto_requerido:N", legend=None), tooltip=["repuesto_requerido", "cantidad"])
        st.altair_chart(chart_repuestos, use_container_width=True)

    st.markdown("### 6. Horas de parada por nivel")
    df_horas_nivel = df_filtrado.groupby("nivel")["tiempo_parada"].sum().reset_index().sort_values("tiempo_parada", ascending=False)
    chart_horas_nivel = alt.Chart(df_horas_nivel).mark_bar().encode(x=alt.X("nivel:N", sort="-y", title="Nivel"), y=alt.Y("tiempo_parada:Q", title="Horas de parada"), color=alt.Color("nivel:N", legend=None), tooltip=["nivel", alt.Tooltip("tiempo_parada:Q", format=".2f")])
    st.altair_chart(chart_horas_nivel, use_container_width=True)

    st.markdown("### 7. Eventos por técnico")
    df_tecnico = df_filtrado.groupby("tecnico").size().reset_index(name="eventos").sort_values("eventos", ascending=False)
    chart_tecnico = alt.Chart(df_tecnico).mark_bar().encode(x=alt.X("eventos:Q", title="Eventos"), y=alt.Y("tecnico:N", sort="-x", title="Técnico"), color=alt.Color("tecnico:N", legend=None), tooltip=["tecnico", "eventos"])
    st.altair_chart(chart_tecnico, use_container_width=True)

    st.markdown("### 8. Horas preventivas vs correctivas")
    df_tipo_mantto = df_filtrado.groupby("tipo_mantenimiento")["tiempo_parada"].sum().reset_index().sort_values("tiempo_parada", ascending=False)
    if df_tipo_mantto.empty:
        st.info("No hay datos de tipo de mantenimiento para graficar.")
    else:
        chart_tipo_mantto = alt.Chart(df_tipo_mantto).mark_bar().encode(x=alt.X("tipo_mantenimiento:N", title="Tipo mantenimiento"), y=alt.Y("tiempo_parada:Q", title="Horas"), color=alt.Color("tipo_mantenimiento:N", legend=None), tooltip=["tipo_mantenimiento", alt.Tooltip("tiempo_parada:Q", format=".2f")])
        st.altair_chart(chart_tipo_mantto, use_container_width=True)

    st.markdown("### 9. Tendencia diaria de eventos")
    df_tendencia = df_filtrado.groupby(df_filtrado["fecha"].dt.date).size().reset_index(name="eventos")
    chart_tendencia = alt.Chart(df_tendencia).mark_line(point=True).encode(x=alt.X("fecha:T", title="Fecha"), y=alt.Y("eventos:Q", title="Eventos"), tooltip=["fecha:T", "eventos"])
    st.altair_chart(chart_tendencia, use_container_width=True)

    st.markdown("---")
    st.subheader("📋 Historial de Eventos")

    columnas_mostrar = ["id", "fecha", "tecnico", "apoyo_1", "apoyo_2", "nivel", "ubicacion", "codigo", "tipo_mantenimiento", "tipo_falla", "causa_preliminar", "repuesto_requerido", "descripcion", "hora_falla", "hora_subsanada", "tiempo_parada", "estado", "foto"]
    df_tabla = df_filtrado[columnas_mostrar].copy()
    df_tabla["fecha"] = df_tabla["fecha"].dt.strftime("%d/%m/%Y")

    st.dataframe(df_tabla, use_container_width=True, hide_index=True)

    st.markdown("---")
    st.subheader("📷 Evidencia fotográfica")

    eventos_con_foto = df_filtrado[df_filtrado["foto"].astype(str).str.contains("base64", na=False)].copy()

    if eventos_con_foto.empty:
        st.info("No hay evidencias fotográficas en los registros filtrados.")
    else:
        eventos_con_foto["selector"] = eventos_con_foto["id"].astype(str) + " | " + eventos_con_foto["codigo"].astype(str) + " | " + eventos_con_foto["tipo_falla"].astype(str) + " | " + eventos_con_foto["fecha"].dt.strftime("%d/%m/%Y")
        seleccion = st.selectbox("Selecciona evento para ver evidencia", eventos_con_foto["selector"].tolist())
        fila = eventos_con_foto[eventos_con_foto["selector"] == seleccion].iloc[0]
        imagen_bytes = obtener_bytes_imagen(fila["foto"])

        if imagen_bytes is None:
            st.warning("No se pudo leer la imagen.")
        else:
            st.image(imagen_bytes, caption=f"Evidencia {fila['id']} - {fila['codigo']}", use_container_width=True)
            st.download_button(label="⬇️ Descargar evidencia JPG", data=imagen_bytes, file_name=f"evidencia_{fila['id']}_{fila['codigo']}.jpg", mime="image/jpeg")

    st.markdown("---")

    csv = df_tabla.to_csv(index=False).encode("utf-8-sig")
    st.download_button(label="⬇️ Descargar historial CSV", data=csv, file_name="historial_eventos_bombeo.csv", mime="text/csv")
